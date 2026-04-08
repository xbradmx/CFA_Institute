"""
Generate DDDS Technical Summary .docx
Run: python generate_technical_summary.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------------------------------------------------------------------------
# COLOURS
# ---------------------------------------------------------------------------
DARK_BLUE  = RGBColor(0x1F, 0x3D, 0x6B)
MID_BLUE   = RGBColor(0x2E, 0x75, 0xB6)
BLACK      = RGBColor(0x00, 0x00, 0x00)
LIGHT_GREY = RGBColor(0xF2, 0xF2, 0xF2)
RED        = RGBColor(0xC0, 0x00, 0x00)


def set_font(run, name="Calibri", size=10, bold=False, italic=False, colour=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if colour:
        run.font.color.rgb = colour


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    if level == 1:
        set_font(run, size=13, bold=True, colour=DARK_BLUE)
    else:
        set_font(run, size=11, bold=True, colour=MID_BLUE)
    return p


def add_body(doc, text, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    set_font(run, size=10, colour=BLACK)
    return p


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.3)
    if bold_prefix:
        r1 = p.add_run(bold_prefix + " ")
        set_font(r1, size=10, bold=True, colour=DARK_BLUE)
        r2 = p.add_run(text)
        set_font(r2, size=10, colour=BLACK)
    else:
        run = p.add_run(text)
        set_font(run, size=10, colour=BLACK)
    return p


def add_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "2E75B6")
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_table_row(table, *cols, header=False):
    row = table.add_row()
    for i, text in enumerate(cols):
        cell = row.cells[i]
        cell.paragraphs[0].clear()
        p = cell.paragraphs[0]
        run = p.add_run(text)
        set_font(run, size=9.5, bold=header, colour=DARK_BLUE if header else BLACK)
        if header:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), "D6E4F0")
            tcPr.append(shd)
    return row


def set_col_width(table, col_idx, width_inches):
    for row in table.rows:
        row.cells[col_idx].width = Inches(width_inches)


def build_doc():
    doc = Document()

    # --- Page margins ---
    section = doc.sections[0]
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

    # =========================================================================
    # TITLE BLOCK
    # =========================================================================
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after  = Pt(2)
    r = title_p.add_run("DDDS — Disclosure Degradation Detection System")
    set_font(r, size=16, bold=True, colour=DARK_BLUE)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.paragraph_format.space_before = Pt(0)
    sub_p.paragraph_format.space_after  = Pt(2)
    r = sub_p.add_run("Technical Summary  |  The Transparency Project")
    set_font(r, size=10.5, italic=True, colour=MID_BLUE)

    inst_p = doc.add_paragraph()
    inst_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    inst_p.paragraph_format.space_before = Pt(0)
    inst_p.paragraph_format.space_after  = Pt(10)
    r = inst_p.add_run("Lancaster University  |  CFA AI Investment Challenge 2026")
    set_font(r, size=9.5, colour=RGBColor(0x60, 0x60, 0x60))

    add_divider(doc)

    # =========================================================================
    # EXECUTIVE SUMMARY
    # =========================================================================
    add_heading(doc, "Executive Summary")

    add_body(doc,
        "Every quarter, public companies publish tens of thousands of pages of regulatory "
        "disclosures. Hidden within these filings is a consistent and exploitable pattern: "
        "before a company's financial condition deteriorates, its disclosures quietly degrade — "
        "specific metrics are replaced with vague language, commitments are dropped, and new "
        "risks appear buried in boilerplate. DDDS is a fully automated system built to detect "
        "this pattern at scale, across an entire sector, before it appears in the financials.",
        space_after=6,
    )

    add_bullet(doc,
        "Corporate disclosure quality degrades measurably in the quarters preceding earnings "
        "shortfalls and credit events. Traditional analyst workflows cannot monitor this signal "
        "systematically across hundreds of companies and filing periods simultaneously.",
        bold_prefix="The Problem."
    )
    add_bullet(doc,
        "DDDS ingests all 10-K and 10-Q filings for 266 US Industrial companies (SIC 3400–3599) "
        "from SEC EDGAR, classifies every risk-factor sentence by topic, scores each sentence for "
        "vagueness and linguistic complexity using fine-tuned FinBERT models, and loads the results "
        "into a Neo4j knowledge graph — enabling temporal, peer-relative, and cross-document "
        "analysis at a scale no individual analyst could replicate.",
        bold_prefix="The Solution."
    )
    add_bullet(doc,
        "The fine-tuned FinBERT vagueness classifier achieved an F1 Macro score of 0.876, and the "
        "complexity classifier achieved 0.826 — both trained on human-labelled SEC disclosure "
        "passages. The two-tier LLM analysis layer (GPT-4o-mini screening followed by Claude Opus "
        "deep analysis) produces structured, evidence-grounded investment signals with direct "
        "citation to source text, and generates formatted two-page investment memos meeting "
        "CFA Standard V(A) auditability requirements.",
        bold_prefix="The Results."
    )
    add_bullet(doc,
        "An analyst using DDDS can query the system for any company in the covered universe and "
        "receive an automatically generated memo identifying specific language deterioration, "
        "peer-relative vagueness anomalies, and management-disclosure contradictions — turning a "
        "multi-day manual review process into a minutes-long automated one. The system is "
        "validated against a Li (2008)-methodology backtesting framework testing whether flagged "
        "companies show lower earnings persistence in the subsequent period.",
        bold_prefix="The Impact."
    )

    add_divider(doc)

    # =========================================================================
    # 1. ARCHITECTURAL DESIGN
    # =========================================================================
    add_heading(doc, "1.  Architectural Design")

    add_body(doc,
        "DDDS is structured as an eleven-stage sequential pipeline with three distinct layers: "
        "a data ingestion layer, a machine learning scoring layer, and a graph-augmented LLM "
        "analysis layer. Each stage produces a persistent artefact consumed by the next, allowing "
        "any stage to be re-run independently without reprocessing upstream steps."
    )

    add_heading(doc, "1.1  Three-Layer Architecture", level=2)

    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    add_table_row(table, "Layer", "Components", header=True)
    rows = [
        ("Data Ingestion",      "SEC EDGAR API collector · HTML/text parser · amendment filter · sentences CSV"),
        ("ML Scoring",          "GPT-4o-mini topic classifier (12 categories) · FinBERT vagueness model · FinBERT complexity model"),
        ("Graph-RAG Analysis",  "Neo4j knowledge graph · GPT-4o-mini screener · Claude Opus analyst · memo generator · heat map · backtester"),
    ]
    for c1, c2 in rows:
        add_table_row(table, c1, c2)
    set_col_width(table, 0, 1.5)
    set_col_width(table, 1, 4.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    add_heading(doc, "1.2  Knowledge Graph Schema", level=2)
    add_body(doc,
        "The central data structure is a Neo4j property graph. Three node types are defined: "
        "Company (identifier, name, sector), Filing (filing_id, filing_type, period), and "
        "RiskFactor (sentence text, topic label, vague_label, vague_prob, complex_label, "
        "complex_prob). Five edge types encode structural and analytical relationships:"
    )
    edge_rows = [
        ("HAS_FILING",    "Company → Filing",           "Ownership"),
        ("HAS_RISK_FACTOR","Filing → RiskFactor",       "Containment"),
        ("NEXT_PERIOD",   "Filing → Filing",             "Temporal succession (same company)"),
        ("PEER_OF",       "Company ↔ Company",           "Same SIC sector peer group"),
        ("SHARES_TOPIC",  "RiskFactor ↔ RiskFactor",    "Same topic, same period, different company"),
    ]
    t2 = doc.add_table(rows=1, cols=3)
    t2.style = "Table Grid"
    add_table_row(t2, "Edge", "Connects", "Purpose", header=True)
    for r in edge_rows:
        row = t2.add_row()
        for i, text in enumerate(r):
            cell = row.cells[i]
            cell.paragraphs[0].clear()
            p2 = cell.paragraphs[0]
            run = p2.add_run(text)
            set_font(run, size=9.5)
    set_col_width(t2, 0, 1.4)
    set_col_width(t2, 1, 1.8)
    set_col_width(t2, 2, 2.8)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    add_divider(doc)

    # =========================================================================
    # 2. WORKFLOW
    # =========================================================================
    add_heading(doc, "2.  Workflow")

    steps = [
        ("Run 0 — Data Collection",
         "The SEC EDGAR full-text search API is queried for all 10-K and 10-Q filings across "
         "266 US Industrial companies. Amended filings (10-K/A, 10-Q/A) are explicitly filtered "
         "to prevent double-counting. All filings are stored as raw HTML."),
        ("Run 1 — Topic Extraction",
         "Each filing's Risk Factors (Item 1A) and MD&A (Item 7/2) sections are isolated using "
         "regex-based section boundary detection on collapsed HTML text. Sentences are extracted "
         "and submitted in batches to the GPT-4o-mini Batch API with a sliding two-sentence "
         "context window for disambiguation. Each sentence receives one of 12 topic labels "
         "(e.g., liquidity_solvency, supply_chain, geopolitical_macro) or 'other'."),
        ("Run 6 — FinBERT Scoring",
         "The fine-tuned vagueness and complexity classifiers are applied to all topic-labelled "
         "sentences, appending vague_label, vague_prob, complex_label, and complex_prob columns "
         "to the master CSV. These sentence-level scores are the core quantitative signal in the system."),
        ("Run 7 — Graph Ingestion",
         "The scored CSV is ingested into Neo4j. Company, Filing, and RiskFactor nodes are "
         "created and all five edge types are established, including automatic NEXT_PERIOD "
         "linking and PEER_OF edges across the sector universe."),
        ("Run 8 — Two-Tier LLM Analysis",
         "Tier 1 (GPT-4o-mini) screens every consecutive filing pair per company per topic, "
         "returning a structured JSON flag (material_change, confidence, change_type). Items "
         "with confidence ≥ 0.70 are escalated to Tier 2 (Claude Opus), which performs "
         "temporal language comparison, peer-relative benchmarking using SHARES_TOPIC graph "
         "queries, and produces a structured JSON finding with signal strength rating "
         "(HIGH / MEDIUM / LOW) and specific analyst action items."),
        ("Runs 9–11 — Output Layer",
         "Run 9 generates formatted two-page investment memos (.docx) from the findings JSON, "
         "including source citations and model provenance consistent with CFA Standard V(A). "
         "Run 10 renders a Bloomberg-style global risk heat map derived from aggregate signal "
         "scores. Run 11 executes a Li (2008) earnings persistence regression to validate that "
         "DDDS-flagged companies exhibit statistically lower forward earnings persistence."),
    ]

    for title, body in steps:
        add_heading(doc, title, level=2)
        add_body(doc, body)

    add_divider(doc)

    # =========================================================================
    # 3. DATA SOURCES
    # =========================================================================
    add_heading(doc, "3.  Data Sources")

    t3 = doc.add_table(rows=1, cols=3)
    t3.style = "Table Grid"
    add_table_row(t3, "Source", "Content", "Volume", header=True)
    data_rows = [
        ("SEC EDGAR API",       "10-K and 10-Q HTML filings",                    "266 companies · SIC 3400–3599"),
        ("OpenAI Batch API",    "GPT-4o-mini topic classification responses",     "~1,064 filings · sentence-level"),
        ("Human annotation",    "Vagueness and complexity labels for FinBERT training", "Training set (stratified by filing type & section)"),
        ("SEC EDGAR XBRL API",  "EPS / ROA earnings data for backtesting",        "Per-company · annual periods"),
        ("yfinance (fallback)",  "Earnings data where XBRL unavailable",          "Supplementary"),
    ]
    for r in data_rows:
        row = t3.add_row()
        for i, text in enumerate(r):
            cell = row.cells[i]
            cell.paragraphs[0].clear()
            run = cell.paragraphs[0].add_run(text)
            set_font(run, size=9.5)
    set_col_width(t3, 0, 1.5)
    set_col_width(t3, 1, 3.0)
    set_col_width(t3, 2, 2.0)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    add_body(doc,
        "All data is sourced from publicly available regulatory filings. No proprietary data "
        "feeds, paid terminals, or non-public information are used. The SEC EDGAR rate limit "
        "(10 requests/second) is respected via a 0.12-second request delay."
    )

    add_divider(doc)

    # =========================================================================
    # 4. DISCUSSION OF RESULTS
    # =========================================================================
    add_heading(doc, "4.  Discussion of Results")

    add_heading(doc, "4.1  FinBERT Model Performance", level=2)
    add_body(doc,
        "Both FinBERT classifiers were fine-tuned on human-labelled SEC disclosure passages, "
        "stratified by filing type (10-K / 10-Q) and section (Risk Factors / MD&A) to avoid "
        "distributional bias. Inter-annotator agreement was measured before training, and a "
        "validation sheet process ensured label consistency across the team."
    )

    t4 = doc.add_table(rows=1, cols=3)
    t4.style = "Table Grid"
    add_table_row(t4, "Model", "Task", "F1 Macro", header=True)
    for r in [("FinBERT — Vagueness", "SPECIFIC / VAGUE", "0.876"),
              ("FinBERT — Complexity", "SIMPLE / COMPLEX",  "0.826")]:
        row = t4.add_row()
        for i, text in enumerate(r):
            cell = row.cells[i]
            cell.paragraphs[0].clear()
            run = cell.paragraphs[0].add_run(text)
            set_font(run, size=9.5)
    set_col_width(t4, 0, 2.0)
    set_col_width(t4, 1, 2.0)
    set_col_width(t4, 2, 2.0)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    add_body(doc,
        "An F1 Macro of 0.876 on the vagueness task is a strong result for financial NLP, "
        "where class imbalance and domain-specific language typically suppress macro scores "
        "relative to accuracy. The complexity classifier's 0.826 reflects the harder task of "
        "distinguishing syntactically dense but informationally rich language from genuine "
        "obfuscation — a distinction that required careful labelling guidelines."
    )

    add_heading(doc, "4.2  Topic Classification", level=2)
    add_body(doc,
        "The GPT-4o-mini topic classifier uses a sliding two-sentence context window to improve "
        "disambiguation at section boundaries. The 12-label taxonomy was designed around "
        "material risk categories for US Industrials, with explicit rules to suppress boilerplate "
        "assignment: approximately 15–25% of sentences are expected to receive the 'other' label. "
        "Temperature is set to 0.0 for reproducibility, and confidence levels (high / medium / low) "
        "are returned alongside each label."
    )

    add_heading(doc, "4.3  LLM Analysis Quality", level=2)
    add_body(doc,
        "The two-tier architecture is specifically designed to control cost while maintaining "
        "analytical depth. GPT-4o-mini handles high-volume screening across all temporal pairs "
        "at low cost, and only items exceeding the 0.70 confidence threshold are escalated to "
        "Claude Opus — which operates within a 180,000-token context limit and receives "
        "structured graph query results rather than raw filings, ensuring analytical focus. "
        "All outputs include exact source text citations, change type classifications, and "
        "structured analyst action items."
    )

    add_divider(doc)

    # =========================================================================
    # 5. MODEL INTERPRETABILITY
    # =========================================================================
    add_heading(doc, "5.  Model Interpretability")

    add_body(doc,
        "Interpretability is a core design requirement given the CFA Standard V(A) auditability "
        "framework the system targets. Several mechanisms are implemented:"
    )

    add_bullet(doc,
        "Every sentence classification by GPT-4o-mini returns a topic_reason field — a one-sentence "
        "natural language explanation of why that label was assigned. This is stored in the CSV "
        "and available for audit.",
        bold_prefix="Topic Classification:"
    )
    add_bullet(doc,
        "Both FinBERT models output a continuous probability score (vague_prob, complex_prob) "
        "in addition to the binary label, enabling downstream threshold tuning and score-based "
        "peer benchmarking rather than binary flagging alone.",
        bold_prefix="FinBERT Scores:"
    )
    add_bullet(doc,
        "The GPT-4o-mini screener returns a structured change_type classification "
        "(new_risk, language_intensification, metric_removal, solvency_language) alongside its "
        "confidence score, making the basis for escalation explicit and auditable.",
        bold_prefix="Screening Transparency:"
    )
    add_bullet(doc,
        "The Claude Opus deep analysis layer is instructed to cite exact language from source "
        "passages in every finding. Analyst memos include model provenance (model name, analysis "
        "date, confidence levels) and a non-reliance disclaimer consistent with CFA standards.",
        bold_prefix="LLM Grounding:"
    )
    add_bullet(doc,
        "The Neo4j graph stores every node and edge that contributed to a finding, allowing "
        "any signal to be traced back to specific filing sentences via the rf_id provenance chain.",
        bold_prefix="Graph Provenance:"
    )

    add_divider(doc)

    # =========================================================================
    # 6. AI USAGE
    # =========================================================================
    add_heading(doc, "6.  AI Usage")

    add_body(doc,
        "DDDS uses AI at three distinct stages, each with a clearly defined and bounded role:"
    )

    t5 = doc.add_table(rows=1, cols=3)
    t5.style = "Table Grid"
    add_table_row(t5, "Model", "Role", "Scope", header=True)
    ai_rows = [
        ("GPT-4o-mini\n(OpenAI Batch API)",
         "Topic classification of risk-factor sentences",
         "12-label taxonomy · one label per sentence · ~1M+ sentences across corpus"),
        ("FinBERT\n(Fine-tuned, local)",
         "Vagueness and complexity scoring",
         "Binary classification with probability output · runs locally on all scored sentences"),
        ("GPT-4o-mini\n(OpenAI Chat API)",
         "Tier 1 screening — temporal change detection",
         "Pairwise filing comparison · structured JSON output · confidence threshold ≥ 0.70"),
        ("Claude Opus 4.6\n(Anthropic API)",
         "Tier 2 deep analysis — temporal, peer, and contradiction analysis",
         "Graph-augmented context · 180k token limit · structured JSON findings"),
    ]
    for r in ai_rows:
        row = t5.add_row()
        for i, text in enumerate(r):
            cell = row.cells[i]
            cell.paragraphs[0].clear()
            run = cell.paragraphs[0].add_run(text)
            set_font(run, size=9)
    set_col_width(t5, 0, 1.5)
    set_col_width(t5, 1, 2.5)
    set_col_width(t5, 2, 3.0)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    add_body(doc,
        "No AI model is given unstructured access to raw data or asked to produce free-form "
        "outputs without a defined schema. All LLM calls use temperature=0.0 (or equivalent) "
        "for reproducibility, structured JSON response formats for parseability, and explicit "
        "system prompts that define scope, output format, and classification rules. "
        "The FinBERT models run locally with no external API dependency, ensuring data "
        "confidentiality for sensitive filing content."
    )

    # =========================================================================
    # FOOTER
    # =========================================================================
    add_divider(doc)
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_p.paragraph_format.space_before = Pt(4)
    r = footer_p.add_run(
        "DDDS — The Transparency Project  ·  Lancaster University  ·  CFA AI Investment Challenge 2026  ·  Confidential"
    )
    set_font(r, size=8, italic=True, colour=RGBColor(0x80, 0x80, 0x80))

    # =========================================================================
    # SAVE
    # =========================================================================
    out_path = "DDDS_Technical_Summary.docx"
    doc.save(out_path)
    print(f"Saved: {out_path}")


if __name__ == "__main__":
    build_doc()
