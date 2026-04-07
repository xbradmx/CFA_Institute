"""
test_topic_extraction.py
========================
Test run: extracts Risk Factors + MD&A from all 5 AAPL filings, sends
sentences to GPT-5-mini for topic classification, and outputs a
colour-coded Word document for human review.

Usage:
    python test_topic_extraction.py

Requirements:
    pip install openai python-docx beautifulsoup4 python-dotenv

Output:
    data/test_topic_highlights.docx
"""

import json
import os
import re
import time
from pathlib import Path

from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from dotenv import load_dotenv
from openai import OpenAI

load_dotenv()

# ---------------------------------------------------------------------------
# CONFIG
# ---------------------------------------------------------------------------
FILINGS_DIR     = Path("data/Filings by company")
OUTPUT_DOCX     = Path("data/test_topic_highlights.docx")
MODEL           = "gpt-5-mini"
TEMPERATURE     = 1
BATCH_SIZE      = 10        # sentences per API call (labelled)
CONTEXT_WINDOW  = 2         # context sentences shown on each side
MAX_RETRIES     = 3
RETRY_DELAY     = 5
# ---------------------------------------------------------------------------


# ---------------------------------------------------------------------------
# TOPIC COLOUR MAP
# ---------------------------------------------------------------------------
TOPIC_COLOURS = {
    "liquidity_solvency":         RGBColor(0xFF, 0x00, 0x00),  # Red
    "debt_financing":             RGBColor(0xDC, 0x14, 0x3C),  # Crimson
    "cost_margin_pressure":       RGBColor(0xFF, 0x8C, 0x00),  # Dark orange
    "supply_chain":               RGBColor(0xFF, 0xD7, 0x00),  # Gold
    "revenue_future_performance": RGBColor(0x00, 0x80, 0x00),  # Green
    "geopolitical_macro":         RGBColor(0xFF, 0xA5, 0x00),  # Orange
    "regulatory_legal":           RGBColor(0x80, 0x00, 0x80),  # Purple
    "operational_product":        RGBColor(0x46, 0x82, 0xB4),  # Steel blue
    "cybersecurity_technology":   RGBColor(0x00, 0x00, 0xCD),  # Medium blue
    "human_capital":              RGBColor(0x20, 0xB2, 0xAA),  # Light sea green
    "financial_sustainability":   RGBColor(0x8B, 0x45, 0x13),  # Saddle brown
    "accounting_audit":           RGBColor(0x8B, 0x00, 0x8B),  # Dark magenta
    "other":                      None,                         # No highlight
}

VALID_LABELS = set(TOPIC_COLOURS.keys())


# ---------------------------------------------------------------------------
# SYSTEM PROMPT
# ---------------------------------------------------------------------------
SYSTEM_PROMPT = """You are an expert financial analyst classifying sentences from SEC filings (US Industrials, SIC 3400-3599) into risk topics.

For each sentence marked [LABEL THIS], assign exactly ONE topic label from below, or 'other' if it doesn't fit any category. Do NOT label sentences marked [CONTEXT ONLY].

TOPIC LABELS:
1. liquidity_solvency - Going concern, cash burn, covenant breach, working capital deficit
2. debt_financing - Leverage, refinancing risk, interest rate exposure, credit rating
3. cost_margin_pressure - Raw material costs, wage inflation, freight costs, margin compression
4. supply_chain - Supplier concentration, component shortages, logistics disruption
5. revenue_future_performance - Revenue decline, backlog reduction, customer concentration, guidance cuts
6. geopolitical_macro - Tariffs, sanctions, FX risk, inflation, recession sensitivity
7. regulatory_legal - Environmental liability, SEC investigation, material litigation, FCPA
8. operational_product - Product recalls, warranty claims, quality failures, plant disruption
9. cybersecurity_technology - Cyber breaches, ransomware, data privacy, IT modernisation risk
10. human_capital - Executive departures, labour shortages, union disputes, succession risk
11. financial_sustainability - Goodwill impairment, pension gaps, dividend sustainability, capex burden
12. accounting_audit - Material weakness, restatements, auditor changes, revenue recognition complexity
other - Generic boilerplate, navigational sentences, factual statements with no risk framing

RULES:
- ONE label per sentence. Choose the most specific applicable topic.
- Going concern language ALWAYS gets liquidity_solvency.
- Generic boilerplate ("we face competition", "results may differ materially") -> other.
- Factual statements with no risk framing -> other.
- Navigational sentences ("see Note 8", "the following table") -> other.
- Financial table data or numeric fragments -> other.
- ~15-25% of sentences are expected to be 'other'.

OUTPUT FORMAT: Respond ONLY with valid JSON:
{
  "results": [
    {"sentence": 5, "label": "cost_margin_pressure", "confidence": "high", "reason": "Discusses raw material cost increases."},
    ...
  ]
}

You MUST return exactly one result for EVERY sentence marked [LABEL THIS]. The count of results must equal the count of [LABEL THIS] sentences. Never skip any."""


# ---------------------------------------------------------------------------
# HTML PARSING
# ---------------------------------------------------------------------------

def extract_text_from_html(path: Path) -> str:
    """Extracts clean text from an EDGAR HTML filing."""
    encodings = ["utf-8", "latin-1", "cp1252"]
    for enc in encodings:
        try:
            with open(path, "r", encoding=enc) as f:
                soup = BeautifulSoup(f.read(), "html.parser")
                for tag in soup(["script", "style", "head"]):
                    tag.decompose()
                return soup.get_text(separator="\n")
        except UnicodeDecodeError:
            continue
    return ""


def clean_text(text: str) -> str:
    """Cleans extracted text: collapses whitespace, removes page numbers."""
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"(?m)^\s*[-\u2013]\s*\d+\s*[-\u2013]\s*$", "", text)
    text = re.sub(r"(?m)^\s*[Pp]age\s+\d+\s*$", "", text)
    text = re.sub(r"[ \t]{3,}", " ", text)
    return text.strip()


# ---------------------------------------------------------------------------
# SECTION EXTRACTION
#
# Strategy:
#   START: Match "Item 1A" / "Item 7" / "Item 2" followed by the known
#          section title. Requires the Item number prefix so we never
#          match boilerplate mentions like "see risk factors".
#
#   END:   Match the FULL heading text of the NEXT section. These are
#          standardised across SEC filings. The full heading text never
#          appears as a cross-reference in prose, only as the actual
#          section heading.
#
#   LONGEST MATCH: If the start pattern matches multiple times (once in
#          the table of contents, once in the actual section), we keep
#          the longest extraction. TOC matches produce tiny strings.
# ---------------------------------------------------------------------------

SECTION_START = {
    "Risk Factors": [
        r"item\s+1a\.?\s*[-\u2013\u2014.\s]*risk\s+factors",
    ],
    "MD&A": [
        r"item\s+7\.?\s*[-\u2013\u2014.\s]*management.s\s+discussion",
        r"item\s+2\.?\s*[-\u2013\u2014.\s]*management.s\s+discussion",
    ],
}

# End patterns use FULL heading text — unique strings that only appear
# as actual section headings, never as cross-references within prose.
SECTION_END = {
    "Risk Factors": {
        "10-K": [
            r"item\s+1b.*unresolved\s+staff\s+comments",
            r"item\s+1c.*cybersecurity",
            r"item\s+2.*properties",
        ],
        "10-Q": [
            r"item\s+2.*unregistered\s+sales",
            r"item\s+2.*management.s\s+discussion",
        ],
    },
    "MD&A": {
        "10-K": [
            r"item\s+7a.*quantitative\s+and\s+qualitative",
        ],
        "10-Q": [
            r"item\s+3.*quantitative\s+and\s+qualitative",
        ],
    },
}


def extract_section(text: str, section_name: str, filing_type: str) -> str | None:
    """
    Extracts a named section from the filing text.

    Uses full heading text for both start and end boundaries to avoid
    false matches on cross-references. Keeps the longest match to skip
    table-of-contents entries.
    """
    start_patterns = SECTION_START.get(section_name, [])
    end_patterns = SECTION_END.get(section_name, {}).get(filing_type, [])

    if not end_patterns:
        print(f"         [!] No end patterns for {section_name} / {filing_type}")
        return None

    lines = text.split("\n")
    best_section = None
    best_length = 0

    for start_pattern in start_patterns:
        for i, line in enumerate(lines):
            if not re.search(start_pattern, line, re.IGNORECASE):
                continue

            start_line = i
            end_line = len(lines)
            found_end = False

            for j in range(i + 5, min(i + 3000, len(lines))):
                line_j = lines[j].strip()
                if len(line_j) < 5:
                    continue
                for end_pattern in end_patterns:
                    if re.search(end_pattern, line_j, re.IGNORECASE):
                        end_line = j
                        found_end = True
                        break
                if found_end:
                    break

            section_text = "\n".join(lines[start_line:end_line]).strip()

            if len(section_text) > best_length:
                best_length = len(section_text)
                best_section = section_text

    if best_section and best_length > 500:
        return best_section
    return None


# ---------------------------------------------------------------------------
# SENTENCE SPLITTING
# ---------------------------------------------------------------------------

def split_into_sentences(text: str) -> list[str]:
    """
    Splits text into individual sentences.
    Filters out short fragments and table data (lines that are mostly numbers).
    """
    protected = text
    abbreviations = [
        "Inc.", "Corp.", "Ltd.", "Co.", "L.P.", "LLC.", "N.A.", "S.A.",
        "Mr.", "Mrs.", "Ms.", "Dr.", "Jr.", "Sr.", "Prof.",
        "No.", "Nos.", "Vol.", "vs.", "etc.", "approx.", "i.e.", "e.g.",
        "U.S.", "U.K.", "E.U.", "ASC.",
    ]
    for abbr in abbreviations:
        protected = protected.replace(abbr, abbr.replace(".", "<<DOT>>"))

    # Protect decimal numbers
    protected = re.sub(r"(\d)\.(\d)", r"\1<<DOT>>\2", protected)

    # Split on sentence-ending punctuation
    raw = re.split(r'(?<=[.!?])\s+(?=[A-Z0-9"\(])', protected)

    # Restore dots
    sentences = [s.replace("<<DOT>>", ".").strip() for s in raw]

    # Filter: min 25 chars and at least 40% alphabetic (skip table data)
    filtered = []
    for s in sentences:
        if len(s) < 25:
            continue
        alpha = sum(1 for c in s if c.isalpha())
        if alpha / max(len(s), 1) < 0.40:
            continue
        filtered.append(s)

    return filtered


# ---------------------------------------------------------------------------
# FILE SELECTION
# ---------------------------------------------------------------------------

def select_test_filings(filings_dir: Path) -> list[dict]:
    """Selects all Apple (AAPL) filings for the test run."""
    aapl_dir = filings_dir / "AAPL"
    if not aapl_dir.exists():
        raise FileNotFoundError(f"AAPL folder not found at {aapl_dir}")

    selected = []
    for filing_type in ["10-K", "10-Q"]:
        type_dir = aapl_dir / filing_type
        if not type_dir.exists():
            continue
        for f in sorted(type_dir.iterdir()):
            if f.suffix.lower() in (".html", ".htm"):
                selected.append({
                    "ticker": "AAPL",
                    "path": f,
                    "filing_type": filing_type,
                })
    return selected


def extract_filing_date(filename: str) -> str:
    """Extracts the filing date from filename like AAON_10-K_2024-02-28_..."""
    match = re.search(r"(\d{4}-\d{2}-\d{2})", filename)
    return match.group(1) if match else "unknown"


# ---------------------------------------------------------------------------
# GPT-5-mini API CALLS
# ---------------------------------------------------------------------------

def build_batches(sentences: list[str]) -> list[dict]:
    """Groups sentences into batches with context overlap."""
    total = len(sentences)
    batches = []

    for start in range(0, total, BATCH_SIZE):
        end = min(start + BATCH_SIZE, total)
        ctx_start = max(0, start - CONTEXT_WINDOW)
        ctx_end = min(total, end + CONTEXT_WINDOW)

        lines = []
        expected_ids = []

        for i in range(ctx_start, ctx_end):
            if i < start or i >= end:
                lines.append(f"[CONTEXT ONLY] Sentence {i}: {sentences[i]}")
            else:
                lines.append(f"[LABEL THIS] Sentence {i}: {sentences[i]}")
                expected_ids.append(i)

        batches.append({
            "sentences_block": "\n\n".join(lines),
            "expected_ids": expected_ids,
        })

    return batches


def call_gpt(client: OpenAI, ticker: str, filing_type: str,
             section: str, sentences_block: str,
             expected_ids: list[int]) -> list[dict] | None:
    """
    Single synchronous API call to GPT-5-mini for topic extraction.
    Retries on missing IDs, accepts partial results on final attempt.
    """
    num_expected = len(expected_ids)
    user_prompt = f"""Company: {ticker}
Filing: {filing_type} -- Section: {section}

Classify each sentence marked [LABEL THIS] into one of the 12 risk topics or 'other'.
Sentences marked [CONTEXT ONLY] are for context only -- do NOT label them.
You MUST return exactly {num_expected} results.

---SENTENCES---
{sentences_block}
---END---"""

    for attempt in range(MAX_RETRIES):
        try:
            response = client.chat.completions.create(
                model=MODEL,
                temperature=TEMPERATURE,
                response_format={"type": "json_object"},
                messages=[
                    {"role": "system", "content": SYSTEM_PROMPT},
                    {"role": "user",   "content": user_prompt},
                ],
            )
            raw = response.choices[0].message.content
            parsed = json.loads(raw)

            if "results" not in parsed:
                raise ValueError("Response missing 'results' key")

            results = parsed["results"]

            # Validate labels
            for r in results:
                if r.get("label") not in VALID_LABELS:
                    r["label"] = "other"

            # Filter to expected IDs only
            expected_set = set(expected_ids)
            results = [r for r in results if r.get("sentence") in expected_set]

            # Check for missing IDs
            returned_ids = {r["sentence"] for r in results}
            missing = expected_set - returned_ids

            if missing and attempt < MAX_RETRIES - 1:
                print(f"\n      [!] Attempt {attempt+1}: missing {len(missing)} IDs, retrying...", end="")
                time.sleep(RETRY_DELAY)
                continue

            if missing:
                print(f"\n      [!] Accepting partial: {len(results)}/{num_expected}", end="")

            return results

        except Exception as e:
            if attempt < MAX_RETRIES - 1:
                print(f"\n      [!] Attempt {attempt+1} failed: {e}", end="")
                time.sleep(RETRY_DELAY * (attempt + 1))
            else:
                print(f"\n      [!] All retries failed: {e}", end="")
                return None


# ---------------------------------------------------------------------------
# WORD DOCUMENT OUTPUT
# ---------------------------------------------------------------------------

def build_highlighted_doc(all_results: list[dict], output_path: Path):
    """Creates a Word doc with colour-coded sentence highlights."""
    doc = Document()

    # Title page
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("DDDS Topic Extraction -- Test Run (AAPL)")
    run.bold = True
    run.font.size = Pt(18)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(
        f"{len(all_results)} filings | Model: {MODEL} | Batch: {BATCH_SIZE}"
    )
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Colour legend
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("Colour Legend")
    run.bold = True
    run.font.size = Pt(14)

    for topic, colour in TOPIC_COLOURS.items():
        if colour is None:
            continue
        p = doc.add_paragraph()
        run = p.add_run(f"  \u25a0 {topic}")
        run.font.color.rgb = colour
        run.font.size = Pt(10)
        run.bold = True

    p = doc.add_paragraph()
    run = p.add_run("  \u25a1 other -- greyed out text, no highlight")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # Each filing
    for filing in all_results:
        doc.add_page_break()

        header = doc.add_paragraph()
        run = header.add_run(
            f"{filing['ticker']} -- {filing['filing_type']} -- {filing['filing_date']}"
        )
        run.bold = True
        run.font.size = Pt(16)

        for section_data in filing["sections"]:
            section_name = section_data["section"]
            labelled = section_data["labelled_sentences"]

            # Section subheader
            sub = doc.add_paragraph()
            run = sub.add_run(section_name)
            run.bold = True
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

            # Stats
            topic_counts = {}
            for s in labelled:
                topic_counts[s["label"]] = topic_counts.get(s["label"], 0) + 1

            non_other = [
                (t, c) for t, c in sorted(topic_counts.items(), key=lambda x: -x[1])
                if t != "other"
            ]
            other_count = topic_counts.get("other", 0)

            stats_parts = [f"{len(labelled)} sentences"]
            if non_other:
                stats_parts.append(", ".join(f"{t}: {c}" for t, c in non_other))
            if other_count:
                stats_parts.append(f"other: {other_count}")

            stats = doc.add_paragraph()
            run = stats.add_run(" | ".join(stats_parts))
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            run.italic = True

            # Sentences
            for s in labelled:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.space_before = Pt(2)

                label = s["label"]
                colour = TOPIC_COLOURS.get(label)

                tag_run = p.add_run(f"[{s['sentence_num']}] ")
                tag_run.font.size = Pt(8)
                tag_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

                text_run = p.add_run(s["text"])
                text_run.font.size = Pt(9)
                if colour is not None:
                    text_run.font.color.rgb = colour
                else:
                    text_run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

                if label != "other":
                    label_run = p.add_run(f"  [{label}]")
                    label_run.font.size = Pt(7)
                    label_run.font.color.rgb = colour
                    label_run.bold = True

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f"\n  Output saved: {output_path}")


# ---------------------------------------------------------------------------
# MAIN
# ---------------------------------------------------------------------------

def main():
    print("\n[DDDS] Topic Extraction -- Test Run (AAPL)")
    print("=" * 55)
    print(f"  Model:      {MODEL}")
    print(f"  Batch size: {BATCH_SIZE}")
    print(f"  Context:    {CONTEXT_WINDOW} sentences each side")

    # Select filings
    print("\n[1/4] Selecting test filings...")
    filings = select_test_filings(FILINGS_DIR)

    if not filings:
        print(f"  [!] No AAPL filings found. Check: {FILINGS_DIR.resolve()}")
        return

    print(f"  Found {len(filings)} AAPL filings:")
    for f in filings:
        print(f"    {f['filing_type']} | {f['path'].name}")

    # Connect
    print("\n[2/4] Connecting to OpenAI...")
    api_key = os.environ.get("OPENAI_API_KEY")
    if not api_key:
        print("  [!] OPENAI_API_KEY not set. Add it to .env")
        return
    client = OpenAI(api_key=api_key)
    print("  Connected")

    # Process
    print("\n[3/4] Processing filings...")
    all_results = []

    for idx, filing_info in enumerate(filings):
        ticker = filing_info["ticker"]
        filing_type = filing_info["filing_type"]
        filepath = filing_info["path"]
        filing_date = extract_filing_date(filepath.name)

        print(f"\n  [{idx+1}/{len(filings)}] {ticker} | {filing_type} | {filing_date}")
        print(f"         File: {filepath.name}")

        raw_text = extract_text_from_html(filepath)
        if not raw_text or len(raw_text) < 500:
            print("         [!] Failed to extract text -- skipping")
            continue

        raw_text = clean_text(raw_text)
        print(f"         Extracted {len(raw_text):,} chars")

        filing_result = {
            "ticker": ticker,
            "filing_type": filing_type,
            "filing_date": filing_date,
            "filename": filepath.name,
            "sections": [],
        }

        for section_name in ["Risk Factors", "MD&A"]:
            section_text = extract_section(raw_text, section_name, filing_type)

            if not section_text:
                print(f"         [!] {section_name} not found -- skipping")
                continue

            print(f"         {section_name}: {len(section_text):,} chars")

            sentences = split_into_sentences(section_text)
            print(f"         -> {len(sentences)} sentences")

            if not sentences:
                continue

            batches = build_batches(sentences)
            all_labelled = []

            for batch_idx, batch in enumerate(batches):
                print(f"           Batch {batch_idx+1}/{len(batches)} "
                      f"({len(batch['expected_ids'])} sent)", end=" ... ")

                results = call_gpt(
                    client=client,
                    ticker=ticker,
                    filing_type=filing_type,
                    section=section_name,
                    sentences_block=batch["sentences_block"],
                    expected_ids=batch["expected_ids"],
                )

                if results:
                    for r in results:
                        r["text"] = sentences[r["sentence"]]
                        r["sentence_num"] = r["sentence"]
                    all_labelled.extend(results)
                    print(f"ok {len(results)}")
                else:
                    print("FAILED")

                time.sleep(0.3)

            filing_result["sections"].append({
                "section": section_name,
                "labelled_sentences": all_labelled,
            })

            if all_labelled:
                topics = {}
                for s in all_labelled:
                    topics[s["label"]] = topics.get(s["label"], 0) + 1
                top_5 = sorted(topics.items(), key=lambda x: -x[1])[:5]
                print(f"         Top: {', '.join(f'{t}: {c}' for t, c in top_5)}")

        all_results.append(filing_result)

    # Build doc
    print(f"\n[4/4] Building highlighted Word document...")
    build_highlighted_doc(all_results, OUTPUT_DOCX)

    # Summary
    total = sum(len(s["labelled_sentences"]) for f in all_results for s in f["sections"])
    other = sum(1 for f in all_results for s in f["sections"]
                for sent in s["labelled_sentences"] if sent["label"] == "other")
    print(f"\n{'='*55}")
    print(f"  Test Run Complete")
    print(f"{'='*55}")
    print(f"  Filings:   {len(all_results)}")
    print(f"  Sentences: {total}")
    print(f"  Other:     {other} ({100*other/max(total,1):.1f}%)")
    print(f"  Output:    {OUTPUT_DOCX}")
    print(f"{'='*55}")


if __name__ == "__main__":
    main()