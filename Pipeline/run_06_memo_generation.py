"""
DDDS - Investment Memo Generator (Run 6b)
==========================================
Reads findings JSON files produced by Run 5 (topic screening) and generates
formatted two-page investment memos as .docx files for each flagged supplier.

Each memo contains:
    - Header: company identity, signal strength, date generated
    - Section 1: Disclosure Quality Trends (quarterly progression)
    - Section 2: Specific Changes with supporting evidence and citations
    - Section 3: Cross-Document Contradictions with parallel quotes
    - Section 4: Peer Context vs sector average
    - Section 5: Historical Pattern Summary
    - Footer: model provenance, confidence levels, non-reliance disclaimer

All outputs include model provenance and confidence levels consistent
with CFA Standard V(A) auditability requirements.

Usage
-----
    # Single company findings file
    python run_06_memo_generation.py --input data/findings/AAPL_findings.json

    # All findings in a directory
    python run_06_memo_generation.py --input-dir data/findings

    # Custom output directory
    python run_06_memo_generation.py --input-dir data/findings --output-dir data/memos
"""

import argparse
import json
import os
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------------------------------------------------------------------------
# CONFIG
# ---------------------------------------------------------------------------
INPUT_DIR  = "data/findings"
OUTPUT_DIR = "data/memos"

# Colours
COLOUR_DARK_BLUE  = RGBColor(0x1F, 0x3D, 0x6B)   # header / section titles
COLOUR_MID_BLUE   = RGBColor(0x2E, 0x75, 0xB6)   # accent lines
COLOUR_RED        = RGBColor(0xC0, 0x00, 0x00)   # HIGH signal
COLOUR_AMBER      = RGBColor(0xED, 0x7D, 0x31)   # MEDIUM signal
COLOUR_GREEN      = RGBColor(0x38, 0x86, 0x38)   # LOW signal
COLOUR_LIGHT_GREY = RGBColor(0xF2, 0xF2, 0xF2)   # table header background
COLOUR_BLACK      = RGBColor(0x00, 0x00, 0x00)

SIGNAL_COLOURS = {
    "HIGH":   COLOUR_RED,
    "MEDIUM": COLOUR_AMBER,
    "LOW":    COLOUR_GREEN,
}

SCREENING_MODEL     = "GPT-4o-mini"
DEEP_ANALYSIS_MODEL = "Claude Opus"
# ---------------------------------------------------------------------------


# ---------------------------------------------------------------------------
# DOCUMENT UTILITIES
# ---------------------------------------------------------------------------

def set_cell_background(cell, hex_colour: str):
    """Sets table cell background colour."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_colour)
    tcPr.append(shd)


def add_horizontal_rule(doc: Document, colour_hex: str = "2E75B6", thickness: int = 12):
    """Adds a coloured horizontal rule as a paragraph border."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(thickness))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), colour_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def add_section_heading(doc: Document, text: str):
    """Adds a styled section heading."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = COLOUR_DARK_BLUE
    run.font.name = "Arial"
    add_horizontal_rule(doc, "2E75B6", 8)
    return p


def add_body_paragraph(doc: Document, text: str, italic: bool = False,
                        indent: bool = False, space_after: int = 6) -> None:
    """Adds a standard body paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.space_before = Pt(2)
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text)
    run.font.size   = Pt(9)
    run.font.name   = "Arial"
    run.font.italic = italic
    run.font.color.rgb = COLOUR_BLACK


def add_quote_block(doc: Document, label: str, quote: str):
    """Adds a labelled quote block with light background indent."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.3)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    label_run = p.add_run(f"{label}: ")
    label_run.bold = True
    label_run.font.size = Pt(8.5)
    label_run.font.name = "Arial"
    label_run.font.color.rgb = COLOUR_DARK_BLUE
    quote_run = p.add_run(f'"{quote}"')
    quote_run.italic = True
    quote_run.font.size = Pt(8.5)
    quote_run.font.name = "Arial"
    quote_run.font.color.rgb = COLOUR_BLACK


def add_key_value_table(doc: Document, rows: list[tuple[str, str]]):
    """Adds a two-column key-value table."""
    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"

    col_widths = [Inches(1.8), Inches(5.2)]

    for i, (key, value) in enumerate(rows):
        row = table.rows[i]
        # Key cell
        key_cell = row.cells[0]
        key_cell.width = col_widths[0]
        set_cell_background(key_cell, "EBF3FB")
        kp = key_cell.paragraphs[0]
        kr = kp.add_run(key)
        kr.bold = True
        kr.font.size = Pt(8.5)
        kr.font.name = "Arial"
        kp.paragraph_format.space_before = Pt(3)
        kp.paragraph_format.space_after  = Pt(3)

        # Value cell
        val_cell = row.cells[1]
        val_cell.width = col_widths[1]
        vp = val_cell.paragraphs[0]
        vr = vp.add_run(str(value) if value else "Not available")
        vr.font.size = Pt(8.5)
        vr.font.name = "Arial"
        vp.paragraph_format.space_before = Pt(3)
        vp.paragraph_format.space_after  = Pt(3)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_disclaimer(doc: Document):
    """Adds the non-reliance disclaimer footer block."""
    add_horizontal_rule(doc, "CCCCCC", 6)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    run = p.add_run(
        "IMPORTANT NOTICE: This document is an automated screening output generated by the "
        "Disclosure Degradation Detection System (DDDS) and is provided for informational "
        "purposes only. It does not constitute investment advice, a recommendation to buy or "
        "sell any security, or a solicitation of any investment decision. All findings are "
        "probabilistic signals requiring independent analyst verification. Recipients should "
        "conduct their own due diligence before making any investment decision. "
        "DDDS outputs are consistent with the role of a screening tool as defined under "
        "CFA Institute Standard V(A) — Diligence and Reasonable Basis."
    )
    run.font.size = Pt(7.5)
    run.font.italic = True
    run.font.name = "Arial"
    run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    p.paragraph_format.space_after = Pt(4)

    p2 = doc.add_paragraph()
    run2 = p2.add_run(
        f"Generated: {datetime.now().strftime('%d %B %Y %H:%M')}  |  "
        f"Screening model: {SCREENING_MODEL}  |  "
        f"Deep analysis model: {DEEP_ANALYSIS_MODEL}  |  "
        f"Source: SEC EDGAR public filings"
    )
    run2.font.size = Pt(7.5)
    run2.font.name = "Arial"
    run2.font.color.rgb = RGBColor(0x60, 0x60, 0x60)


# ---------------------------------------------------------------------------
# MEMO SECTIONS
# ---------------------------------------------------------------------------

def build_header(doc: Document, company: dict, signal_strength: str, flags_count: int):
    """Builds the memo header with company identity and signal badge."""

    # Title row
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after  = Pt(2)
    title_run = p_title.add_run("SUPPLIER DISCLOSURE RISK MEMO")
    title_run.bold = True
    title_run.font.size = Pt(13)
    title_run.font.name = "Arial"
    title_run.font.color.rgb = COLOUR_DARK_BLUE

    # Company name
    p_company = doc.add_paragraph()
    p_company.paragraph_format.space_after = Pt(2)
    comp_run = p_company.add_run(
        f"{company.get('name', 'Unknown Company')} ({company.get('identifier', '')})"
    )
    comp_run.bold = True
    comp_run.font.size = Pt(11)
    comp_run.font.name = "Arial"
    comp_run.font.color.rgb = COLOUR_DARK_BLUE

    # Signal strength badge
    p_signal = doc.add_paragraph()
    p_signal.paragraph_format.space_after = Pt(2)
    badge_run = p_signal.add_run(f"SIGNAL STRENGTH: {signal_strength}")
    badge_run.bold = True
    badge_run.font.size = Pt(10)
    badge_run.font.name = "Arial"
    badge_run.font.color.rgb = SIGNAL_COLOURS.get(signal_strength, COLOUR_BLACK)

    # Meta row
    p_meta = doc.add_paragraph()
    p_meta.paragraph_format.space_after = Pt(6)
    meta_run = p_meta.add_run(
        f"Sector: {company.get('sector', 'US Industrials')}  |  "
        f"Flagged categories: {flags_count}  |  "
        f"Generated: {datetime.now().strftime('%d %B %Y')}"
    )
    meta_run.font.size = Pt(8.5)
    meta_run.font.name = "Arial"
    meta_run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

    add_horizontal_rule(doc, "1F3D6B", 16)


def build_section_1_trends(doc: Document, flagged_items: list):
    """Section 1: Disclosure Quality Trends - quarterly progression."""
    add_section_heading(doc, "1. Disclosure Quality Trends")

    if not flagged_items:
        add_body_paragraph(doc, "No material trends identified in this period.")
        return

    # Summary table of all flagged categories
    rows = [("Risk Category", "Period", "Change Type", "Confidence")]
    for item in flagged_items:
        rows.append((
            item.get("risk_category", ""),
            f"{item.get('prev_period', '?')} → {item.get('curr_period', '?')}",
            item.get("change_type", "").replace("_", " ").title(),
            f"{item.get('confidence', 0):.0%}",
        ))

    table = doc.add_table(rows=len(rows), cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    col_widths = [Inches(2.0), Inches(1.6), Inches(2.2), Inches(1.2)]

    for i, row_data in enumerate(rows):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            cell.width = col_widths[j]
            p = cell.paragraphs[0]
            r = p.add_run(cell_text)
            r.font.size = Pt(8.5)
            r.font.name = "Arial"
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(3)
            if i == 0:
                r.bold = True
                set_cell_background(cell, "1F3D6B")
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            else:
                set_cell_background(cell, "F2F2F2" if i % 2 == 0 else "FFFFFF")

    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def build_section_2_changes(doc: Document, temporal_analysis: dict, flagged_items: list):
    """Section 2: Specific Changes with supporting evidence."""
    add_section_heading(doc, "2. Specific Changes & Supporting Evidence")

    if not temporal_analysis:
        add_body_paragraph(doc, "No temporal analysis available.")
        return

    # Deterioration assessment
    assessment = temporal_analysis.get("deterioration_assessment", "")
    if assessment:
        add_body_paragraph(doc, assessment)

    # Metric removal flag
    if temporal_analysis.get("metric_removal"):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Inches(0.3)
        flag_run = p.add_run("⚠ METRIC REMOVAL DETECTED: ")
        flag_run.bold = True
        flag_run.font.size = Pt(9)
        flag_run.font.name = "Arial"
        flag_run.font.color.rgb = COLOUR_RED
        detail_run = p.add_run(temporal_analysis.get("metric_removal_detail", ""))
        detail_run.font.size = Pt(9)
        detail_run.font.name = "Arial"

    # Changes list
    changes = temporal_analysis.get("changes_detected", [])
    if changes:
        add_body_paragraph(doc, "Detected changes:", space_after=3)
        for change in changes:
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent  = Inches(0.4)
            p.paragraph_format.space_after  = Pt(3)
            r = p.add_run(str(change))
            r.font.size = Pt(8.5)
            r.font.name = "Arial"

    # Supporting evidence from flagged items
    if flagged_items:
        add_body_paragraph(doc, "Document evidence:", space_after=3)
        for item in flagged_items[:3]:  # cap at 3 for page constraint
            cat = item.get("risk_category", "").title()
            prev_p = item.get("prev_period", "")
            curr_p = item.get("curr_period", "")
            reason = item.get("screening_reason", "")

            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after  = Pt(2)
            p.paragraph_format.left_indent  = Inches(0.2)
            cat_run = p.add_run(f"{cat} ({prev_p} → {curr_p}): ")
            cat_run.bold = True
            cat_run.font.size = Pt(8.5)
            cat_run.font.name = "Arial"
            reason_run = p.add_run(reason)
            reason_run.font.size = Pt(8.5)
            reason_run.font.name = "Arial"
            reason_run.font.italic = True

            # Truncated passage comparison
            prev_text = item.get("prev_text", "")[:300]
            curr_text = item.get("curr_text", "")[:300]
            if prev_text:
                add_quote_block(doc, f"Previous ({prev_p})", prev_text + "...")
            if curr_text:
                add_quote_block(doc, f"Current ({curr_p})", curr_text + "...")


def build_section_3_peer(doc: Document, peer_analysis: dict):
    """Section 3: Peer context vs sector average."""
    add_section_heading(doc, "3. Peer Context & Sector Benchmarking")

    if not peer_analysis:
        add_body_paragraph(doc, "No peer analysis available.")
        return

    # Vagueness deviation
    deviation = peer_analysis.get("vagueness_deviation", "")
    deviation_map = {
        "above_average": ("ABOVE AVERAGE vagueness", COLOUR_RED),
        "average":       ("AVERAGE vagueness", COLOUR_AMBER),
        "below_average": ("BELOW AVERAGE vagueness", COLOUR_GREEN),
    }
    if deviation in deviation_map:
        label, colour = deviation_map[deviation]
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        dev_run = p.add_run(f"Disclosure vagueness: {label} relative to sector peers")
        dev_run.bold = True
        dev_run.font.size = Pt(9)
        dev_run.font.name = "Arial"
        dev_run.font.color.rgb = colour

    # Peer comparison narrative
    comparison = peer_analysis.get("peer_comparison", "")
    if comparison:
        add_body_paragraph(doc, comparison)

    # Omitted risks
    omitted = peer_analysis.get("omitted_risks", [])
    if omitted:
        add_body_paragraph(doc, "Risks disclosed by peers but absent from this filing:", space_after=3)
        for risk in omitted:
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Inches(0.4)
            p.paragraph_format.space_after  = Pt(3)
            r = p.add_run(str(risk))
            r.font.size  = Pt(8.5)
            r.font.name  = "Arial"
            r.font.color.rgb = COLOUR_RED


def build_section_4_signal(doc: Document, overall_signal: dict):
    """Section 4: Overall signal and analyst actions."""
    add_section_heading(doc, "4. Overall Assessment & Analyst Actions")

    if not overall_signal:
        add_body_paragraph(doc, "No overall signal assessment available.")
        return

    signal_strength = overall_signal.get("signal_strength", "UNKNOWN")
    summary = overall_signal.get("summary", "")
    actions = overall_signal.get("analyst_actions", [])

    # Signal box
    add_key_value_table(doc, [
        ("Signal Strength", signal_strength),
        ("Summary",         summary),
    ])

    # Analyst actions
    if actions:
        add_body_paragraph(doc, "Recommended analyst actions:", space_after=3)
        for i, action in enumerate(actions, 1):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent  = Inches(0.3)
            p.paragraph_format.space_after  = Pt(3)
            num_run = p.add_run(f"{i}. ")
            num_run.bold = True
            num_run.font.size = Pt(9)
            num_run.font.name = "Arial"
            act_run = p.add_run(str(action))
            act_run.font.size = Pt(9)
            act_run.font.name = "Arial"

    # Model provenance
    add_horizontal_rule(doc, "CCCCCC", 4)
    prov_rows = [
        ("Tier 1 screening",   f"{SCREENING_MODEL} — material change detection"),
        ("Tier 2 analysis",    f"{DEEP_ANALYSIS_MODEL} — temporal and peer analysis"),
        ("Data source",        "SEC EDGAR public filings (10-K, 10-Q)"),
        ("Confidence basis",   "Consensus dual-call GPT-4o-mini screening (threshold ≥ 0.70)"),
    ]
    add_key_value_table(doc, prov_rows)


# ---------------------------------------------------------------------------
# DOCUMENT SETUP
# ---------------------------------------------------------------------------

def configure_document(doc: Document):
    """Sets page margins and default font."""
    section = doc.sections[0]
    section.page_width   = Inches(8.5)
    section.page_height  = Inches(11)
    section.top_margin    = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin   = Inches(1.0)
    section.right_margin  = Inches(1.0)

    # Default paragraph style
    style = doc.styles["Normal"]
    font  = style.font
    font.name = "Arial"
    font.size = Pt(9)


# ---------------------------------------------------------------------------
# MAIN GENERATOR
# ---------------------------------------------------------------------------

def generate_memo(findings: dict, output_path: str):
    """
    Generates a single investment memo .docx from a findings dict.
    """
    company         = findings.get("company", {})
    flags_count     = findings.get("flags_count", 0)
    flagged_items   = findings.get("flagged_items", [])
    deep_analysis   = findings.get("deep_analysis", {})

    temporal        = deep_analysis.get("temporal_analysis", {})
    peer            = deep_analysis.get("peer_analysis", {})
    overall_signal  = deep_analysis.get("overall_signal", {})
    signal_strength = overall_signal.get("signal_strength", "UNKNOWN")

    doc = Document()
    configure_document(doc)

    # Build each section
    build_header(doc, company, signal_strength, flags_count)
    build_section_1_trends(doc, flagged_items)
    build_section_2_changes(doc, temporal, flagged_items)
    build_section_3_peer(doc, peer)
    build_section_4_signal(doc, overall_signal)
    add_disclaimer(doc)

    os.makedirs(Path(output_path).parent, exist_ok=True)
    doc.save(output_path)


def load_findings(path: str) -> dict | None:
    """Loads and validates a findings JSON file."""
    try:
        with open(path, "r") as f:
            findings = json.load(f)
        if "company" not in findings or "deep_analysis" not in findings:
            print(f"  [!] {path}: missing required keys, skipping")
            return None
        return findings
    except Exception as e:
        print(f"  [!] Failed to load {path}: {e}")
        return None


def main():
    parser = argparse.ArgumentParser(description="DDDS Investment Memo Generator")
    parser.add_argument("--input",      default=None,       help="Single findings JSON file")
    parser.add_argument("--input-dir",  default=INPUT_DIR,  help="Directory of findings JSON files")
    parser.add_argument("--output-dir", default=OUTPUT_DIR, help="Directory for output .docx memos")
    args = parser.parse_args()

    print(f"\n[DDDS] Investment Memo Generator")
    print(f"  Output dir: {args.output_dir}\n")

    os.makedirs(args.output_dir, exist_ok=True)

    # Collect files to process
    if args.input:
        files = [args.input]
    else:
        files = sorted(Path(args.input_dir).glob("*_findings.json"))
        files = [str(f) for f in files]

    if not files:
        print(f"  [!] No findings JSON files found in '{args.input_dir}'")
        print(f"  Make sure Run 8 (graph_rag.py) has been executed first.")
        return

    print(f"  Processing {len(files)} findings file(s)...\n")

    generated = []
    skipped   = []

    for path in files:
        filename = Path(path).name
        print(f"  [{len(generated)+len(skipped)+1}/{len(files)}] {filename}", end=" ... ")

        findings = load_findings(path)
        if not findings:
            skipped.append(filename)
            continue

        identifier = findings.get("company", {}).get("identifier", "unknown")
        signal     = findings.get("deep_analysis", {}).get(
            "overall_signal", {}
        ).get("signal_strength", "UNKNOWN")

        out_filename = f"{identifier}_investment_memo.docx"
        out_path     = os.path.join(args.output_dir, out_filename)

        try:
            generate_memo(findings, out_path)
            generated.append((identifier, signal, out_path))
            print(f"✓ {signal} signal  →  {out_filename}")
        except Exception as e:
            print(f"ERROR: {e}")
            skipped.append(filename)

    # Summary
    print(f"\n{'='*55}")
    print(f"  Memo Generation Summary")
    print(f"{'='*55}")
    print(f"  Generated:  {len(generated)} memos")
    print(f"  Skipped:    {len(skipped)}")

    if generated:
        high   = [i for i, s, _ in generated if s == "HIGH"]
        medium = [i for i, s, _ in generated if s == "MEDIUM"]
        low    = [i for i, s, _ in generated if s == "LOW"]
        print(f"\n  By signal strength:")
        print(f"    HIGH:    {len(high)}  {high[:5]}")
        print(f"    MEDIUM:  {len(medium)}  {medium[:5]}")
        print(f"    LOW:     {len(low)}")
        print(f"\n  Memos saved to: {args.output_dir}")

    print(f"{'='*55}")
    print(f"\n[Done]")


if __name__ == "__main__":
    main()