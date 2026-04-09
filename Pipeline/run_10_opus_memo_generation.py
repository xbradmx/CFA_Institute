"""
DDDS - Investment Memo Generation (Run 10)
============================================
Reads Opus Pass 1 analysis JSONs and FinBERT prediction scores,
calls Claude Opus 4.6 to generate structured narrative sections,
and produces formatted two-page .docx investment memos with inline
filing citations and dual signal assessment (FinBERT + Claude).

Usage
-----
    python "run_10_-_memo_generation.py" --all
    python "run_10_-_memo_generation.py" --ticker AAPL
    python "run_10_-_memo_generation.py" --all --use-cached
    python "run_10_-_memo_generation.py" --all --dry-run
"""

import argparse
import csv
import io
import json
import os
import random
import re
import time
from collections import defaultdict
from datetime import datetime
from html.parser import HTMLParser
from pathlib import Path

import anthropic
from neo4j import GraphDatabase
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from dotenv import load_dotenv

load_dotenv()

# ---------------------------------------------------------------------------
# CONFIG
# ---------------------------------------------------------------------------
OPUS_ANALYSIS_DIR   = "data/Graph Rag Creation Data/opus_analysis"
ANALYST_OUTPUTS_DIR = "data/analyst_outputs"
FILINGS_BASE_DIR    = "data/Filings by company"
OUTPUT_DIR          = "data/analyst_outputs/memos"
MEMO_CACHE_DIR      = "data/analyst_outputs/memo_cache"

OPUS_MODEL  = "claude-opus-4-6"
MAX_TOKENS  = 1500
RETRY_DELAY = 60

NEO4J_URI      = os.environ.get("NEO4J_URI", "neo4j+s://2e14ef0f.databases.neo4j.io")
NEO4J_USER     = os.environ.get("NEO4J_USER", "2e14ef0f")
NEO4J_PASSWORD = os.environ.get("NEO4J_PASSWORD")

# Colours
COLOUR_DARK_BLUE  = RGBColor(0x1F, 0x3D, 0x6B)
COLOUR_MID_BLUE   = RGBColor(0x2E, 0x75, 0xB6)
COLOUR_RED        = RGBColor(0xC0, 0x00, 0x00)
COLOUR_AMBER      = RGBColor(0xED, 0x7D, 0x31)
COLOUR_GREEN      = RGBColor(0x38, 0x86, 0x38)
COLOUR_BLACK      = RGBColor(0x00, 0x00, 0x00)
COLOUR_GREY       = RGBColor(0x60, 0x60, 0x60)

SIGNAL_COLOURS = {"HIGH": COLOUR_RED, "MEDIUM": COLOUR_AMBER, "LOW": COLOUR_GREEN}

_ACCESSION_RE = re.compile(r'\d{10}-\d{2}-\d{6}')
QUOTE_PATTERN = re.compile(r"\[QUOTE:([^\]]+)\]")


# ---------------------------------------------------------------------------
# NEO4J
# ---------------------------------------------------------------------------

class GraphQuerier:
    def __init__(self, uri: str, user: str, password: str):
        self.driver = GraphDatabase.driver(uri, auth=(user, password))

    def close(self):
        self.driver.close()

    def get_blocks_batch(self, block_ids: list[str]) -> dict[str, str]:
        query = """
        UNWIND $ids AS bid
        MATCH (d:DisclosureBlock {block_id: bid})
        RETURN d.block_id AS block_id, d.text AS text
        """
        with self.driver.session() as session:
            results = session.run(query, ids=block_ids)
            return {r["block_id"]: r["text"] for r in results}

    def get_company_metadata(self, ticker: str) -> dict | None:
        query = """
        MATCH (c:Company {ticker: $ticker})
        RETURN c.ticker AS ticker, c.name AS name, c.sector AS sector
        """
        with self.driver.session() as session:
            result = session.run(query, ticker=ticker)
            record = result.single()
            return dict(record) if record else None

    def get_all_company_names(self, tickers: list[str]) -> dict[str, str]:
        """Batch fetch company names. Returns {ticker: name}."""
        query = """
        UNWIND $tickers AS t
        MATCH (c:Company {ticker: t})
        RETURN c.ticker AS ticker, c.name AS name
        """
        with self.driver.session() as session:
            results = session.run(query, tickers=tickers)
            return {r["ticker"]: r["name"] for r in results}


def collect_block_ids(flags: list[dict]) -> list[str]:
    ids = set()
    for flag in flags:
        block_ids = flag.get("block_ids") or {}
        if not isinstance(block_ids, dict):
            continue
        for bid in block_ids.values():
            if bid:
                ids.add(bid)
    return list(ids)


def build_block_lookup(block_texts: dict[str, str]) -> dict[str, dict]:
    return {
        bid: {"text": text, "section": "", "vague_prob": None,
              "vague_label": "", "complex_prob": None, "complex_label": ""}
        for bid, text in block_texts.items()
    }


# ---------------------------------------------------------------------------
# SENTENCE / CSV LOOKUP
# ---------------------------------------------------------------------------

def load_block_lookup(ticker: str, scores_dir: str) -> dict[str, dict]:
    lookup: dict[str, dict] = {}
    pattern = re.compile(rf"^{re.escape(ticker)}_.*_ddds_scores\.csv$", re.IGNORECASE)

    if not os.path.isdir(scores_dir):
        return lookup

    for filename in os.listdir(scores_dir):
        if not pattern.match(filename):
            continue
        filepath = os.path.join(scores_dir, filename)
        try:
            with open(filepath, "r", encoding="utf-8-sig") as fh:
                raw = fh.read()
            marker = "SENTENCE LEVEL DETAIL"
            idx = raw.find(marker)
            if idx == -1:
                continue
            csv_text = raw[idx + len(marker):].strip()
            reader = csv.DictReader(io.StringIO(csv_text))
            for row in reader:
                sid = row.get("sentence_id", "").strip()
                if not sid:
                    continue
                lookup[sid] = {
                    "text": row.get("text", ""),
                    "section": row.get("section", ""),
                    "vague_prob": _safe_float(row.get("vague_prob")),
                    "vague_label": row.get("vague_label", ""),
                    "complex_prob": _safe_float(row.get("complex_prob")),
                    "complex_label": row.get("complex_label", ""),
                    "demeaned_vague": _safe_float(row.get("demeaned_vague")),
                    "demeaned_complex": _safe_float(row.get("demeaned_complex")),
                    "outlier": row.get("outlier", "").strip().upper() == "TRUE",
                }
        except Exception as e:
            print(f"    [!] Error reading {filename}: {e}")
    return lookup


def _safe_float(val) -> float | None:
    try:
        return float(val)
    except (TypeError, ValueError):
        return None


def _filing_key(sentence_id: str) -> str:
    m = _ACCESSION_RE.search(sentence_id)
    if m:
        return sentence_id[:m.end()]
    parts = sentence_id.rsplit("_", 2)
    return parts[0] if len(parts) >= 3 else sentence_id


def _ticker_from_filename(filename: str) -> str | None:
    parts = filename.split("_")
    return parts[0] if len(parts) >= 2 else None


def compute_filing_scores(lookup: dict[str, dict]) -> dict:
    if not lookup:
        return {"filings": [], "latest": None, "previous": None, "delta": None,
                "total_sentences": 0, "num_filings": 0}

    by_filing: dict[str, list] = defaultdict(list)
    for sid, s in lookup.items():
        by_filing[_filing_key(sid)].append(s)

    filings = []
    for key, sentences in by_filing.items():
        vague_probs = [s["vague_prob"] for s in sentences if s["vague_prob"] is not None]
        complex_probs = [s["complex_prob"] for s in sentences if s["complex_prob"] is not None]
        key_parts = key.split("_")
        date_str = key_parts[2] if len(key_parts) > 2 else "unknown"
        filing_type = key_parts[1] if len(key_parts) > 1 else "?"
        filings.append({
            "filing_key": key, "date": date_str, "filing_type": filing_type,
            "vagueness_avg": round(sum(vague_probs) / len(vague_probs), 4) if vague_probs else None,
            "complexity_avg": round(sum(complex_probs) / len(complex_probs), 4) if complex_probs else None,
            "sentence_count": len(sentences),
        })

    filings.sort(key=lambda x: x["date"])
    latest = filings[-1] if filings else None
    previous = filings[-2] if len(filings) >= 2 else None

    delta = None
    if latest and previous:
        dv = (round(latest["vagueness_avg"] - previous["vagueness_avg"], 4)
              if latest["vagueness_avg"] is not None and previous["vagueness_avg"] is not None
              else None)
        dc = (round(latest["complexity_avg"] - previous["complexity_avg"], 4)
              if latest["complexity_avg"] is not None and previous["complexity_avg"] is not None
              else None)
        delta = {"vagueness": dv, "complexity": dc}

    return {
        "filings": filings, "latest": latest, "previous": previous,
        "delta": delta, "total_sentences": len(lookup), "num_filings": len(filings),
    }


# ---------------------------------------------------------------------------
# UNIVERSE STATS
# ---------------------------------------------------------------------------

def compute_universe_stats(scores_dir: str) -> dict:
    company_scores: dict[str, list[float]] = defaultdict(list)
    company_complex: dict[str, list[float]] = defaultdict(list)

    if not os.path.isdir(scores_dir):
        return {"mean_vague": None, "std_vague": None, "mean_complex": None,
                "std_complex": None, "p25_vague": None, "p75_vague": None,
                "num_companies": 0}

    for filename in os.listdir(scores_dir):
        if not filename.endswith("_ddds_scores.csv"):
            continue
        ticker = _ticker_from_filename(filename)
        if not ticker:
            continue
        filepath = os.path.join(scores_dir, filename)
        try:
            with open(filepath, "r", encoding="utf-8-sig") as fh:
                raw = fh.read()
            marker = "SENTENCE LEVEL DETAIL"
            idx = raw.find(marker)
            if idx == -1:
                continue
            csv_text = raw[idx + len(marker):].strip()
            reader = csv.DictReader(io.StringIO(csv_text))
            for row in reader:
                vp = _safe_float(row.get("vague_prob"))
                cp = _safe_float(row.get("complex_prob"))
                if vp is not None:
                    company_scores[ticker].append(vp)
                if cp is not None:
                    company_complex[ticker].append(cp)
        except Exception:
            continue

    vague_avgs = sorted([sum(p) / len(p) for p in company_scores.values() if p])
    complex_avgs = sorted([sum(p) / len(p) for p in company_complex.values() if p])

    if not vague_avgs:
        return {"mean_vague": None, "std_vague": None, "mean_complex": None,
                "std_complex": None, "p25_vague": None, "p75_vague": None,
                "num_companies": 0}

    import statistics
    n = len(vague_avgs)
    return {
        "mean_vague": statistics.mean(vague_avgs),
        "std_vague": statistics.stdev(vague_avgs) if n > 1 else 0,
        "mean_complex": statistics.mean(complex_avgs) if complex_avgs else None,
        "std_complex": statistics.stdev(complex_avgs) if len(complex_avgs) > 1 else 0,
        "p25_vague": vague_avgs[int(n * 0.25)],
        "p75_vague": vague_avgs[int(n * 0.75)],
        "p25_complex": complex_avgs[int(len(complex_avgs) * 0.25)] if complex_avgs else None,
        "p75_complex": complex_avgs[int(len(complex_avgs) * 0.75)] if complex_avgs else None,
        "num_companies": n,
    }


def classify_finbert_signal(company_vagueness: float | None,
                            universe: dict) -> str:
    if company_vagueness is None or universe.get("p75_vague") is None:
        return "UNKNOWN"
    if company_vagueness >= universe["p75_vague"]:
        return "HIGH"
    elif company_vagueness <= universe["p25_vague"]:
        return "LOW"
    else:
        return "MEDIUM"


# ---------------------------------------------------------------------------
# PAGE ESTIMATION — search HTML files by sentence text
# ---------------------------------------------------------------------------

class _TagStripper(HTMLParser):
    def __init__(self):
        super().__init__()
        self.text = []
    def handle_data(self, d):
        self.text.append(d)

# Cache: {html_path: plain_text}
_html_cache: dict[str, str] = {}

# Cache: {ticker: [(filing_type, date, html_path), ...]}
_ticker_filings_cache: dict[str, list[tuple[str, str, str]]] = {}


def _get_ticker_filings(ticker: str, filings_base: str) -> list[tuple[str, str, str]]:
    """Returns [(filing_type, date, html_path), ...] for a ticker, sorted by date desc."""
    if ticker in _ticker_filings_cache:
        return _ticker_filings_cache[ticker]

    results = []
    ticker_dir = os.path.join(filings_base, ticker)
    if not os.path.isdir(ticker_dir):
        _ticker_filings_cache[ticker] = []
        return []

    for filing_type_dir in os.listdir(ticker_dir):
        type_path = os.path.join(ticker_dir, filing_type_dir)
        if not os.path.isdir(type_path):
            continue
        for fname in os.listdir(type_path):
            if not fname.endswith(".html"):
                continue
            # Parse: TICKER_TYPE_DATE_ACCESSION.html
            parts = fname.replace(".html", "").split("_")
            if len(parts) >= 3:
                filing_type = parts[1]
                date = parts[2]
                results.append((filing_type, date, os.path.join(type_path, fname)))

    results.sort(key=lambda x: x[1], reverse=True)
    _ticker_filings_cache[ticker] = results
    return results


def _get_plain_text(html_path: str) -> str:
    """Returns cached plain text for an HTML file."""
    if html_path not in _html_cache:
        try:
            with open(html_path, "r", encoding="utf-8", errors="ignore") as f:
                raw = f.read()
            stripper = _TagStripper()
            stripper.feed(raw)
            _html_cache[html_path] = "".join(stripper.text)
        except Exception:
            _html_cache[html_path] = ""
    return _html_cache[html_path]


def find_sentence_in_filings(ticker: str, sentence_text: str,
                             filings_base: str,
                             chars_per_page: int = 4000) -> str | None:
    """
    Searches all HTML filings for a ticker to find the sentence.
    Returns a citation string like 'see 10-K filed 2024-08-02, p.15'
    or None if not found.
    """
    if not sentence_text or len(sentence_text) < 20:
        return None

    search_text = sentence_text[:100]
    filings = _get_ticker_filings(ticker, filings_base)

    for filing_type, date, html_path in filings:
        plain = _get_plain_text(html_path)
        if not plain:
            continue

        idx = plain.find(search_text)
        if idx != -1:
            page = (idx // chars_per_page) + 1
            return f"see {filing_type} filed {date}, p.{page}"

    return None


# ---------------------------------------------------------------------------
# RESOLVE BLOCK_IDS
# ---------------------------------------------------------------------------

def resolve_flag_sentences(flags: list[dict], lookup: dict[str, dict]) -> list[dict]:
    enriched = []
    for flag in flags:
        block_ids = flag.get("block_ids") or {}
        if not isinstance(block_ids, dict):
            enriched.append({**flag, "resolved_sentences": {}})
            continue
        resolved = {}
        for role, bid in block_ids.items():
            if bid and bid in lookup:
                s = lookup[bid]
                resolved[role] = {
                    "sentence_id": bid, "text": s["text"],
                    "section": s["section"], "vague_prob": s["vague_prob"],
                    "vague_label": s["vague_label"],
                    "complex_prob": s["complex_prob"],
                    "complex_label": s["complex_label"],
                }
            elif bid:
                resolved[role] = {"sentence_id": bid, "text": "[not found]"}
        enriched.append({**flag, "resolved_sentences": resolved})
    return enriched


# ---------------------------------------------------------------------------
# OPUS MEMO PROMPT
# ---------------------------------------------------------------------------

MEMO_SYSTEM_PROMPT = """\
You are a senior buy-side equity analyst writing a concise two-page investment \
memo for a portfolio manager. You are part of the Disclosure Degradation \
Detection System (DDDS), a supply chain monitoring tool that detects when \
upstream suppliers become less transparent in their SEC filings.

You will receive:
1. The Opus Pass 1 flag analysis with resolved sentence text
2. The overall company signal assessment (Claude Analysis Signal)
3. The FinBERT statistical signal (sector-relative vagueness ranking)
4. Per-filing FinBERT model scores with trend
5. The company name and ticker

IMPORTANT RULES:
- When the two signals diverge (e.g. FinBERT HIGH but Claude LOW), discuss \
both perspectives and explain the divergence. Neither signal overrides the other.
- NEVER recommend that analysts "request the DDDS team" or "ask the system" \
to do anything. If there is a limitation (e.g. poor peer selection), note it \
as a known limitation — do not frame it as a task for someone to fix.
- When discussing specific findings, reference the sentence_id so the system \
can generate a page citation. Format as [QUOTE:sentence_id]. Use a MAXIMUM \
of 4 citations. Only reference sentence_ids from the context.
- Always refer to the company by its full name on first mention, then by \
ticker thereafter.

THE ENTIRE MEMO MUST FIT ON TWO PRINTED PAGES. Be ruthlessly concise. \
Professional prose for a CFA charterholder audience. No bullet points.

Respond ONLY with valid JSON:
{
  "executive_summary": "ONE paragraph, 60-80 words. Lead with supply chain \
risk implication. State BOTH signal strengths (Claude Analysis and FinBERT). \
If they diverge, note it. End with recommended action.",

  "disclosure_quality_assessment": "ONE paragraph, 50-70 words. Interpret \
the FinBERT scores relative to sector. Reference trend direction and \
the company's percentile position.",

  "key_findings": "TWO paragraphs, 100-150 words total. The 2-3 most \
material findings. Include [QUOTE:sentence_id] for key evidence. Do not \
enumerate every flag.",

  "peer_comparison": "ONE paragraph, 40-60 words. How does disclosure \
compare to peers? Note any known peer selection limitations rather than \
recommending fixes.",

  "recommended_actions": "ONE paragraph, 50-70 words. 2-3 concrete analyst \
actions framed as supply chain due diligence steps the analyst themselves \
can take."
}"""


def build_memo_context(ticker: str,
                       company_name: str,
                       opus_analysis: dict,
                       enriched_flags: list[dict],
                       agg_scores: dict,
                       finbert_signal: str,
                       universe_stats: dict) -> str:
    sections = []

    overall = opus_analysis.get("overall", {})
    claude_signal = overall.get("signal_strength", "UNKNOWN")
    if company_name != ticker:
        sections.append(f"COMPANY: {company_name} ({ticker})")
    else:
        sections.append(f"COMPANY: {ticker}")
    sections.append(f"CLAUDE ANALYSIS SIGNAL: {claude_signal}")
    sections.append(f"FINBERT SIGNAL: {finbert_signal}")
    if claude_signal != finbert_signal:
        sections.append(f"NOTE: Signals diverge. Claude assessed disclosure "
                        f"changes as {claude_signal}; FinBERT ranks this company's "
                        f"vagueness as {finbert_signal} relative to the sector.")
    sections.append("")

    # FinBERT scores per filing
    sections.append("=" * 50)
    sections.append("FINBERT DISCLOSURE QUALITY SCORES")
    sections.append("=" * 50)

    for filing in agg_scores.get("filings", []):
        line = f"  {filing['filing_type']} filed {filing['date']}"
        if filing["vagueness_avg"] is not None:
            line += f" | Vagueness: {filing['vagueness_avg']:.4f}"
        if filing["complexity_avg"] is not None:
            line += f" | Complexity: {filing['complexity_avg']:.4f}"
        line += f" | {filing['sentence_count']} sentences"
        sections.append(line)

    delta = agg_scores.get("delta")
    if delta:
        def _fmt(v):
            if v is None: return "N/A"
            return f"{'+' if v >= 0 else ''}{v:.4f}"
        sections.append(f"Change (latest vs previous): Vagueness {_fmt(delta.get('vagueness'))} | "
                        f"Complexity {_fmt(delta.get('complexity'))}")

    if universe_stats.get("mean_vague") is not None:
        latest = agg_scores.get("latest")
        sections.append(f"\nSector universe ({universe_stats['num_companies']} companies):")
        sections.append(f"  Mean vagueness: {universe_stats['mean_vague']:.4f}")
        sections.append(f"  P25: {universe_stats['p25_vague']:.4f} | "
                        f"P75: {universe_stats['p75_vague']:.4f}")
        if latest and latest["vagueness_avg"] is not None:
            diff = latest["vagueness_avg"] - universe_stats["mean_vague"]
            sections.append(f"  This company vs mean: {'+' if diff >= 0 else ''}{diff:.4f}")
    sections.append("")

    # Overall assessment
    sections.append("=" * 50)
    sections.append("OPUS PASS 1 — OVERALL ASSESSMENT")
    sections.append("=" * 50)
    sections.append(f"Signal: {claude_signal} | "
                    f"Genuine: {overall.get('genuine_count', '?')} | "
                    f"Borderline: {overall.get('borderline_count', '?')} | "
                    f"Dismissed: {overall.get('dismissed_count', '?')}")
    sections.append(f"Summary: {overall.get('summary', 'N/A')}")
    for i, c in enumerate(overall.get("top_concerns", []), 1):
        sections.append(f"  {i}. {c}")
    sections.append("")

    # Active flags
    active = [f for f in enriched_flags if f.get("assessment") in ("GENUINE", "BORDERLINE")]
    if active:
        sections.append("=" * 50)
        sections.append("ACTIVE FLAGS WITH RESOLVED SENTENCES")
        sections.append("=" * 50)
        for flag in active:
            sections.append(f"\n--- Flag {flag.get('flag_number', '?')}: "
                            f"{flag.get('topic', '?')} [{flag.get('layer', '?')}] "
                            f"— {flag.get('assessment', '?')} ---")
            sections.append(f"Reasoning: {flag.get('reasoning', 'N/A')}")
            if flag.get("investigation_action"):
                sections.append(f"Investigation: {flag['investigation_action']}")
            for role, info in flag.get("resolved_sentences", {}).items():
                sid = info.get("sentence_id", "?")
                text = info.get("text", "[not found]")
                vp = info.get("vague_prob")
                sections.append(f"  [{role.upper()}] {sid}")
                if vp is not None:
                    sections.append(f"    Vague: {vp:.4f} ({info.get('vague_label', '')}) | "
                                    f"Complex: {info.get('complex_prob', '?')} "
                                    f"({info.get('complex_label', '')})")
                sections.append(f"    \"{text[:300]}\"")
            sections.append("")

    return "\n".join(sections)


# ---------------------------------------------------------------------------
# OPUS MEMO CALL
# ---------------------------------------------------------------------------

def call_opus_memo(client: anthropic.Anthropic,
                   context: str,
                   ticker: str) -> dict | None:
    non_rate_failures = 0
    rate_limit_retries = 0
    MAX_RATE_RETRIES = 10

    while True:
        try:
            response = client.messages.create(
                model=OPUS_MODEL,
                max_tokens=MAX_TOKENS,
                system=MEMO_SYSTEM_PROMPT,
                messages=[{"role": "user", "content": context}]
            )
            raw = response.content[0].text
            raw = raw.replace("```json", "").replace("```", "").strip()
            try:
                result = json.loads(raw)
            except json.JSONDecodeError:
                print(f"\n      [!] {ticker}: JSON parse failed")
                return None
            result["_meta"] = {
                "ticker": ticker, "model": OPUS_MODEL,
                "input_tokens": response.usage.input_tokens,
                "output_tokens": response.usage.output_tokens,
                "timestamp": datetime.now().isoformat(),
            }
            return result
        except Exception as e:
            err_str = str(e).lower()
            is_rate_limit = ("429" in str(e) or "rate" in err_str or
                             "overloaded" in err_str or "capacity" in err_str)
            if is_rate_limit:
                rate_limit_retries += 1
                if rate_limit_retries >= MAX_RATE_RETRIES:
                    print(f"\n      [!] {ticker}: {MAX_RATE_RETRIES} rate limit retries, giving up.")
                    return None
                wait = RETRY_DELAY + random.uniform(1, 10)
                print(f"\n      [!] {ticker}: Rate limited ({rate_limit_retries}/"
                      f"{MAX_RATE_RETRIES}), retrying in {wait:.0f}s...", flush=True)
                time.sleep(wait)
            else:
                non_rate_failures += 1
                if non_rate_failures >= 3:
                    print(f"\n      [!] {ticker}: 3 failures, giving up. Last: {e}")
                    return None
                time.sleep(15 * non_rate_failures)


# ---------------------------------------------------------------------------
# CACHING
# ---------------------------------------------------------------------------

def get_memo_cache_path(ticker: str, cache_dir: str) -> str:
    return os.path.join(cache_dir, f"{ticker}_memo.json")

def load_memo_cache(ticker: str, cache_dir: str) -> dict | None:
    path = get_memo_cache_path(ticker, cache_dir)
    if os.path.exists(path):
        with open(path, "r") as f:
            return json.load(f)
    return None

def save_memo_cache(ticker: str, data: dict, cache_dir: str) -> None:
    os.makedirs(cache_dir, exist_ok=True)
    with open(get_memo_cache_path(ticker, cache_dir), "w") as f:
        json.dump(data, f, indent=2)


# ---------------------------------------------------------------------------
# DOCX UTILITIES
# ---------------------------------------------------------------------------

def set_cell_background(cell, hex_colour: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_colour)
    tcPr.append(shd)


def add_horizontal_rule(doc: Document, colour_hex: str = "2E75B6", thickness: int = 12):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(thickness))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), colour_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_section_heading(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = COLOUR_DARK_BLUE
    run.font.name = "Arial"
    add_horizontal_rule(doc, "2E75B6", 8)


def add_body_paragraph(doc: Document, text: str, italic: bool = False,
                       space_after: int = 6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.name = "Arial"
    run.font.italic = italic
    run.font.color.rgb = COLOUR_BLACK


def add_key_value_table(doc: Document, rows: list[tuple[str, str]]):
    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    col_widths = [Inches(2.0), Inches(5.0)]
    for i, (key, value) in enumerate(rows):
        row = table.rows[i]
        key_cell = row.cells[0]
        key_cell.width = col_widths[0]
        set_cell_background(key_cell, "EBF3FB")
        kp = key_cell.paragraphs[0]
        kr = kp.add_run(key)
        kr.bold = True
        kr.font.size = Pt(8.5)
        kr.font.name = "Arial"
        kp.paragraph_format.space_before = Pt(3)
        kp.paragraph_format.space_after = Pt(3)
        val_cell = row.cells[1]
        val_cell.width = col_widths[1]
        vp = val_cell.paragraphs[0]
        vr = vp.add_run(str(value) if value else "Not available")
        vr.font.size = Pt(8.5)
        vr.font.name = "Arial"
        vp.paragraph_format.space_before = Pt(3)
        vp.paragraph_format.space_after = Pt(3)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


# ---------------------------------------------------------------------------
# NARRATIVE RENDERING — inline page citations, suppress broken ones
# ---------------------------------------------------------------------------

def render_narrative_with_citations(doc: Document, narrative_text: str,
                                    lookup: dict[str, dict],
                                    ticker: str,
                                    filings_base: str = FILINGS_BASE_DIR):
    if not narrative_text:
        add_body_paragraph(doc, "Not available.", italic=True)
        return

    def _replace_quote(match):
        sid = match.group(1).strip()
        if sid not in lookup:
            return ""
        text = lookup[sid].get("text", "")
        citation = find_sentence_in_filings(ticker, text, filings_base)
        if citation:
            return f"({citation})"
        else:
            return ""  # suppress if we can't find it

    for para_text in narrative_text.split("\n\n"):
        para_text = para_text.strip()
        if not para_text:
            continue
        display_text = QUOTE_PATTERN.sub(_replace_quote, para_text)
        # Clean up double spaces from suppressed citations
        display_text = re.sub(r"  +", " ", display_text).strip()
        if display_text:
            add_body_paragraph(doc, display_text)


# ---------------------------------------------------------------------------
# DOCX BUILDER
# ---------------------------------------------------------------------------

def configure_document(doc: Document):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(9)


def build_memo_docx(ticker: str,
                    company_name: str,
                    opus_analysis: dict,
                    memo_narrative: dict,
                    agg_scores: dict,
                    block_lookup: dict[str, dict],
                    finbert_signal: str,
                    universe_stats: dict,
                    output_path: str,
                    filings_base: str = FILINGS_BASE_DIR):

    overall = opus_analysis.get("overall", {})
    claude_signal = overall.get("signal_strength", "UNKNOWN")
    genuine_count = overall.get("genuine_count", 0)
    borderline_count = overall.get("borderline_count", 0)
    dismissed_count = overall.get("dismissed_count", 0)

    doc = Document()
    configure_document(doc)

    # ===== HEADER =====
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(2)
    title_run = p_title.add_run("SUPPLIER DISCLOSURE RISK MEMO")
    title_run.bold = True
    title_run.font.size = Pt(14)
    title_run.font.name = "Arial"
    title_run.font.color.rgb = COLOUR_DARK_BLUE

    # Full company name + ticker
    p_company = doc.add_paragraph()
    p_company.paragraph_format.space_after = Pt(4)
    if company_name != ticker:
        comp_run = p_company.add_run(f"{company_name} ({ticker})")
    else:
        comp_run = p_company.add_run(ticker)
    comp_run.bold = True
    comp_run.font.size = Pt(12)
    comp_run.font.name = "Arial"
    comp_run.font.color.rgb = COLOUR_DARK_BLUE

    # DUAL SIGNAL
    p_signals = doc.add_paragraph()
    p_signals.paragraph_format.space_after = Pt(2)

    run_label1 = p_signals.add_run("CLAUDE ANALYSIS: ")
    run_label1.bold = True
    run_label1.font.size = Pt(10)
    run_label1.font.name = "Arial"
    run_label1.font.color.rgb = COLOUR_DARK_BLUE
    run_val1 = p_signals.add_run(claude_signal)
    run_val1.bold = True
    run_val1.font.size = Pt(10)
    run_val1.font.name = "Arial"
    run_val1.font.color.rgb = SIGNAL_COLOURS.get(claude_signal, COLOUR_BLACK)

    run_sep = p_signals.add_run("    |    ")
    run_sep.font.size = Pt(10)
    run_sep.font.name = "Arial"
    run_sep.font.color.rgb = COLOUR_GREY

    run_label2 = p_signals.add_run("FINBERT VAGUENESS: ")
    run_label2.bold = True
    run_label2.font.size = Pt(10)
    run_label2.font.name = "Arial"
    run_label2.font.color.rgb = COLOUR_DARK_BLUE
    run_val2 = p_signals.add_run(finbert_signal)
    run_val2.bold = True
    run_val2.font.size = Pt(10)
    run_val2.font.name = "Arial"
    run_val2.font.color.rgb = SIGNAL_COLOURS.get(finbert_signal, COLOUR_BLACK)

    # Meta line
    p_meta = doc.add_paragraph()
    p_meta.paragraph_format.space_after = Pt(4)
    meta_run = p_meta.add_run(
        f"Generated: {datetime.now().strftime('%d %B %Y')}  |  "
        f"Genuine: {genuine_count}  |  Borderline: {borderline_count}  |  "
        f"Dismissed: {dismissed_count}"
    )
    meta_run.font.size = Pt(8.5)
    meta_run.font.name = "Arial"
    meta_run.font.color.rgb = COLOUR_GREY

    add_horizontal_rule(doc, "1F3D6B", 16)

    # ===== FINBERT SCORE TABLE =====
    score_rows = []
    latest = agg_scores.get("latest")
    previous = agg_scores.get("previous")

    if latest:
        parts = []
        if latest["vagueness_avg"] is not None:
            parts.append(f"Vagueness: {latest['vagueness_avg']:.4f}")
        if latest["complexity_avg"] is not None:
            parts.append(f"Complexity: {latest['complexity_avg']:.4f}")
        score_rows.append((f"Most Recent ({latest['filing_type']} {latest['date']})",
                           "  |  ".join(parts) if parts else "N/A"))

    if previous:
        parts = []
        if previous["vagueness_avg"] is not None:
            parts.append(f"Vagueness: {previous['vagueness_avg']:.4f}")
        if previous["complexity_avg"] is not None:
            parts.append(f"Complexity: {previous['complexity_avg']:.4f}")
        score_rows.append((f"Prior ({previous['filing_type']} {previous['date']})",
                           "  |  ".join(parts) if parts else "N/A"))

    def _fmt(v):
        if v is None: return "N/A"
        return f"{'+' if v >= 0 else ''}{v:.4f}"

    delta = agg_scores.get("delta")
    if delta:
        score_rows.append(("Change (latest vs prior)",
                           f"Vagueness: {_fmt(delta.get('vagueness'))}  |  "
                           f"Complexity: {_fmt(delta.get('complexity'))}"))

    if universe_stats.get("mean_vague") is not None:
        score_rows.append(("Sector Average",
                           f"Vagueness: {universe_stats['mean_vague']:.4f}  |  "
                           f"P25: {universe_stats['p25_vague']:.4f}  |  "
                           f"P75: {universe_stats['p75_vague']:.4f}  "
                           f"({universe_stats['num_companies']} companies)"))

    score_rows.append(("Model Performance",
                       "Vagueness F1: 0.876  |  Complexity F1: 0.826"))

    add_key_value_table(doc, score_rows)

    # ===== SECTIONS =====
    add_section_heading(doc, "1. Executive Summary")
    render_narrative_with_citations(doc, memo_narrative.get("executive_summary", ""),
                                    block_lookup, ticker, filings_base)

    add_section_heading(doc, "2. Disclosure Quality Assessment")
    render_narrative_with_citations(doc, memo_narrative.get("disclosure_quality_assessment", ""),
                                    block_lookup, ticker, filings_base)

    add_section_heading(doc, "3. Key Findings & Evidence")
    render_narrative_with_citations(doc, memo_narrative.get("key_findings", ""),
                                    block_lookup, ticker, filings_base)

    add_section_heading(doc, "4. Peer Comparison")
    render_narrative_with_citations(doc, memo_narrative.get("peer_comparison", ""),
                                    block_lookup, ticker, filings_base)

    add_section_heading(doc, "5. Recommended Analyst Actions")
    render_narrative_with_citations(doc, memo_narrative.get("recommended_actions", ""),
                                    block_lookup, ticker, filings_base)

    # ===== PROVENANCE =====
    add_horizontal_rule(doc, "CCCCCC", 4)
    prov_rows = [
        ("FinBERT classifiers", "Vagueness (F1 0.876) + Complexity (F1 0.826)"),
        ("Screening", "GPT-4o-mini | Deep analysis: Claude Opus 4.6"),
        ("Data source", "SEC EDGAR (10-K, 10-Q) | Graph: Neo4j"),
    ]
    add_key_value_table(doc, prov_rows)

    # ===== DISCLAIMER =====
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    run = p.add_run(
        "IMPORTANT NOTICE: This is an automated screening output from the Disclosure "
        "Degradation Detection System (DDDS). It does not constitute investment advice. "
        "All findings are probabilistic signals requiring independent verification. "
        "Consistent with CFA Institute Standard V(A) \u2014 Diligence and Reasonable Basis."
    )
    run.font.size = Pt(7)
    run.font.italic = True
    run.font.name = "Arial"
    run.font.color.rgb = COLOUR_GREY

    p2 = doc.add_paragraph()
    run2 = p2.add_run(
        f"Generated: {datetime.now().strftime('%d %B %Y %H:%M')}  |  "
        f"DDDS v1.0 \u2014 The Transparency Project"
    )
    run2.font.size = Pt(7)
    run2.font.name = "Arial"
    run2.font.color.rgb = COLOUR_GREY

    os.makedirs(Path(output_path).parent, exist_ok=True)
    doc.save(output_path)


# ---------------------------------------------------------------------------
# MAIN
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(description="DDDS Investment Memo Generation (Run 10)")
    parser.add_argument("--ticker", default=None)
    parser.add_argument("--all", action="store_true")
    parser.add_argument("--test", nargs="?", const=1, type=int, default=None)
    parser.add_argument("--opus-dir", default=OPUS_ANALYSIS_DIR)
    parser.add_argument("--scores-dir", default=ANALYST_OUTPUTS_DIR)
    parser.add_argument("--filings-dir", default=FILINGS_BASE_DIR)
    parser.add_argument("--output-dir", default=OUTPUT_DIR)
    parser.add_argument("--cache-dir", default=MEMO_CACHE_DIR)
    parser.add_argument("--use-cached", action="store_true")
    parser.add_argument("--dry-run", action="store_true")
    parser.add_argument("--neo4j-uri", default=NEO4J_URI)
    parser.add_argument("--neo4j-user", default=NEO4J_USER)
    parser.add_argument("--neo4j-password", default=NEO4J_PASSWORD)
    args = parser.parse_args()

    if not args.ticker and not args.all and args.test is None:
        print("Specify --ticker TICKER, --all, or --test [N]")
        return

    print(f"\n[DDDS] Investment Memo Generation \u2014 Run 10")
    print(f"  Opus dir:    {args.opus_dir}")
    print(f"  Scores dir:  {args.scores_dir}")
    print(f"  Filings dir: {args.filings_dir}")
    print(f"  Output dir:  {args.output_dir}")
    print(f"  Model:       {OPUS_MODEL}")

    # Discover Opus analysis files
    opus_files = sorted(Path(args.opus_dir).glob("*_opus_analysis.json"))
    ticker_map = {f.stem.replace("_opus_analysis", ""): str(f) for f in opus_files}
    print(f"  Found {len(ticker_map)} Opus analysis files")

    if not ticker_map:
        print("  [!] No Opus analysis files found. Run 9 first.")
        return

    # Universe stats
    print(f"  Computing universe stats...", end=" ", flush=True)
    universe_stats = compute_universe_stats(args.scores_dir)
    if universe_stats["num_companies"] > 0:
        print(f"{universe_stats['num_companies']} companies | "
              f"mean vague: {universe_stats['mean_vague']:.4f} | "
              f"P25: {universe_stats['p25_vague']:.4f} | "
              f"P75: {universe_stats['p75_vague']:.4f}")
    else:
        print("no data found")

    # Targets
    if args.ticker:
        if args.ticker not in ticker_map:
            print(f"  [!] No Opus analysis for {args.ticker}")
            return
        targets = [args.ticker]
    elif args.test is not None:
        n = max(1, args.test)
        targets = sorted(random.sample(list(ticker_map.keys()),
                                        min(n, len(ticker_map))))
        print(f"  Test: {', '.join(targets)}")
    else:
        targets = sorted(ticker_map.keys())

    client = None
    if not args.dry_run:
        api_key = os.environ.get("ANTHROPIC_API_KEY")
        if not api_key:
            print("  [!] ANTHROPIC_API_KEY not set. Aborting.")
            return
        client = anthropic.Anthropic(api_key=api_key)

    # Neo4j
    neo4j_password = args.neo4j_password or os.environ.get("NEO4J_PASSWORD")
    graph = None
    if neo4j_password:
        try:
            graph = GraphQuerier(args.neo4j_uri, args.neo4j_user, neo4j_password)
            print(f"  Neo4j: connected")
        except Exception as e:
            print(f"  [!] Neo4j failed: {e}")
    else:
        print("  [!] NEO4J_PASSWORD not set")

    # Batch-fetch company names from Neo4j
    company_names: dict[str, str] = {}
    if graph:
        print(f"  Fetching company names...", end=" ", flush=True)
        try:
            company_names = graph.get_all_company_names(targets)
            print(f"{len(company_names)} from Neo4j", end="", flush=True)
        except Exception as e:
            print(f"Neo4j failed ({e})", end="", flush=True)

    # Fallback: load from company_filing_summary.xlsx for any missing names
    summary_path = os.path.join("data", "company_filing_summary.xlsx")
    if os.path.exists(summary_path):
        try:
            import openpyxl
            wb = openpyxl.load_workbook(summary_path, read_only=True)
            ws = wb.active
            header = [cell.value for cell in next(ws.iter_rows(min_row=1, max_row=1))]
            ticker_col = header.index("Ticker") if "Ticker" in header else None
            name_col = header.index("Company Name") if "Company Name" in header else None
            if ticker_col is not None and name_col is not None:
                filled = 0
                for row in ws.iter_rows(min_row=2, values_only=True):
                    t = row[ticker_col]
                    n = row[name_col]
                    if t and n and t in targets and t not in company_names:
                        company_names[t] = n
                        filled += 1
                if filled:
                    print(f" + {filled} from xlsx", end="", flush=True)
            wb.close()
        except Exception as e:
            print(f" xlsx fallback failed ({e})", end="", flush=True)
    print()

    print()
    os.makedirs(args.output_dir, exist_ok=True)
    os.makedirs(args.cache_dir, exist_ok=True)

    results = {"generated": [], "skipped": [], "failed": []}
    total_in = 0
    total_out = 0

    for i, ticker in enumerate(targets, 1):
        company_name = company_names.get(ticker) or ticker
        print(f"  [{i}/{len(targets)}] {company_name} ({ticker})", end=" ... ", flush=True)

        try:
            with open(ticker_map[ticker], "r") as f:
                opus_analysis = json.load(f)
        except Exception as e:
            print(f"SKIP ({e})")
            results["skipped"].append(ticker)
            continue

        claude_signal = opus_analysis.get("overall", {}).get("signal_strength", "?")

        # CSV lookup
        csv_lookup = load_block_lookup(ticker, args.scores_dir)
        agg_scores = compute_filing_scores(csv_lookup)
        print(f"{len(csv_lookup)} sent", end=" ... ", flush=True)

        # FinBERT signal
        latest = agg_scores.get("latest")
        company_vagueness = latest["vagueness_avg"] if latest else None
        finbert_signal = classify_finbert_signal(company_vagueness, universe_stats)

        # Neo4j lookup
        flags = opus_analysis.get("flags", [])
        neo4j_lookup: dict[str, dict] = {}
        if graph:
            block_ids = collect_block_ids(flags)
            if block_ids:
                block_texts = graph.get_blocks_batch(block_ids)
                neo4j_lookup = build_block_lookup(block_texts)

        block_lookup = {**neo4j_lookup, **csv_lookup}

        # Resolve flags
        enriched_flags = resolve_flag_sentences(flags, block_lookup)
        resolved_count = sum(
            1 for f in enriched_flags
            for r in f.get("resolved_sentences", {}).values()
            if r.get("text") and "[not found" not in r["text"]
        )
        print(f"{resolved_count} resolved | C:{claude_signal} F:{finbert_signal}",
              end=" ... ", flush=True)

        # Cache
        memo_narrative = None
        if args.use_cached:
            memo_narrative = load_memo_cache(ticker, args.cache_dir)
            if memo_narrative:
                print("cached", end=" ... ", flush=True)

        if args.dry_run:
            context = build_memo_context(ticker, company_name, opus_analysis,
                                          enriched_flags, agg_scores,
                                          finbert_signal, universe_stats)
            print(f"dry-run | ~{len(context) // 4:,} tokens")
            continue

        # Call Opus
        if memo_narrative is None:
            context = build_memo_context(ticker, company_name, opus_analysis,
                                          enriched_flags, agg_scores,
                                          finbert_signal, universe_stats)
            print("calling Opus", end=" ... ", flush=True)
            memo_narrative = call_opus_memo(client, context, ticker)

            if memo_narrative is None:
                print("FAILED")
                results["failed"].append(ticker)
                continue

            save_memo_cache(ticker, memo_narrative, args.cache_dir)
            total_in += memo_narrative.get("_meta", {}).get("input_tokens", 0)
            total_out += memo_narrative.get("_meta", {}).get("output_tokens", 0)

        # Build docx
        out_path = os.path.join(args.output_dir, f"{ticker}_investment_memo.docx")
        try:
            build_memo_docx(ticker, company_name, opus_analysis, memo_narrative,
                            agg_scores, block_lookup, finbert_signal,
                            universe_stats, out_path, args.filings_dir)
            print("done")
            results["generated"].append((ticker, claude_signal, finbert_signal))
        except Exception as e:
            print(f"DOCX ERROR: {e}")
            results["failed"].append(ticker)

    # Summary
    print(f"\n{'=' * 55}")
    print(f"  Memo Generation Summary")
    print(f"{'=' * 55}")
    print(f"  Generated:  {len(results['generated'])}")
    print(f"  Skipped:    {len(results['skipped'])}")
    print(f"  Failed:     {len(results['failed'])}")

    if results["generated"]:
        c_high = [t for t, cs, fs in results["generated"] if cs == "HIGH"]
        c_med = [t for t, cs, fs in results["generated"] if cs == "MEDIUM"]
        c_low = [t for t, cs, fs in results["generated"] if cs == "LOW"]
        print(f"\n  Claude Signal:  HIGH {len(c_high)} | MEDIUM {len(c_med)} | LOW {len(c_low)}")

        f_high = [t for t, cs, fs in results["generated"] if fs == "HIGH"]
        f_med = [t for t, cs, fs in results["generated"] if fs == "MEDIUM"]
        f_low = [t for t, cs, fs in results["generated"] if fs == "LOW"]
        print(f"  FinBERT Signal: HIGH {len(f_high)} | MEDIUM {len(f_med)} | LOW {len(f_low)}")

        divergent = [(t, cs, fs) for t, cs, fs in results["generated"] if cs != fs]
        if divergent:
            print(f"\n  Divergent signals ({len(divergent)}):")
            for t, cs, fs in divergent[:15]:
                print(f"    {t}: Claude={cs}, FinBERT={fs}")

    if total_in > 0:
        cost = (total_in * 15 / 1_000_000) + (total_out * 75 / 1_000_000)
        print(f"\n  Tokens: {total_in:,} in + {total_out:,} out")
        print(f"  Est. cost: ${cost:.2f}")

    if results["failed"]:
        print(f"\n  Failed: {', '.join(results['failed'])}")

    if graph:
        graph.close()
    _html_cache.clear()
    _ticker_filings_cache.clear()

    print(f"\n  Memos: {args.output_dir}")
    print(f"{'=' * 55}\n[Done]")


if __name__ == "__main__":
    main()